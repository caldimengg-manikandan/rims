from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, BackgroundTasks, Form, Request, Query
from fastapi.responses import ORJSONResponse
from app.core.timezone import get_ist_now, to_naive_ist
from sqlalchemy import or_, and_, func, text, extract, inspect as sa_inspect
from sqlalchemy.orm import Session, joinedload, selectinload, load_only
from app.core.storage import upload_file, get_signed_url, get_public_url
import os
import json
import logging
import time
from datetime import datetime, timezone, timedelta
from app.infrastructure.database import get_db, SessionLocal
from app.domain.models import User, Application, Job, ResumeExtraction, Interview, InterviewReport, ResumeExtractionVersion
from app.domain.schemas import (
    ApplicationCreate,
    ApplicationStatusUpdate,
    ApplicationResponse,
    ApplicationDetailResponse,
    ApplicationSummaryResponse,
    ApplicationNotesUpdate,
    JobSummary,
    ResumeExtractionSummary,
    InterviewSummary,
    InterviewReportSummary,
    HasAppliedResponse,
    ApplicationListResponse,
    BulkDeleteEmailsRequest,
)
from app.core.auth import get_current_hr, get_current_admin
from app.core.ownership import validate_hr_ownership
from app.services.ai_service import parse_resume_with_ai, extract_basic_candidate_info
from app.services.email_service import send_application_received_email, send_rejected_email, send_approved_for_interview_email
from passlib.context import CryptContext
from sqlalchemy.exc import IntegrityError
import re
import hashlib
import asyncio
 
from typing import Optional, List

logger = logging.getLogger(__name__)

# Semaphore: allows only ONE concurrent AI resume analysis.
# This prevents Groq/AI rate-limit errors when many resumes are uploaded in a batch.
_ai_analysis_semaphore = asyncio.Semaphore(1)

# Persisted hint for HR (marker stripped in API responses). Heuristic also sets extraction_degraded.
RIMS_EXTRACTION_DEGRADED_MARKER = "[[rims:extraction_degraded]]"


def _strip_extraction_marker(notes: Optional[str]) -> Optional[str]:
    if not notes:
        return notes
    stripped = notes.replace(RIMS_EXTRACTION_DEGRADED_MARKER, "").strip()
    return stripped if stripped else None


def _append_extraction_degraded_marker(application: Application) -> None:
    current = application.hr_notes or ""
    if RIMS_EXTRACTION_DEGRADED_MARKER in current:
        return
    application.hr_notes = (current.rstrip() + "\n" + RIMS_EXTRACTION_DEGRADED_MARKER).strip()


def _heuristic_extraction_degraded(application: Application) -> bool:
    re = application.resume_extraction
    if not re:
        return False
    summary = re.summary or ""
    if "AI was unable to generate a summary for this resume." in summary:
        return True
    try:
        skills = json.loads(re.extracted_skills or "[]")
    except Exception:
        return False
    if skills == ["General Profile"]:
        return True
    # Avoid triggering a lazy-load of extracted_text (it can be large).
    # If extracted_text isn't loaded in the current query, we can't reliably apply this heuristic.
    try:
        insp = sa_inspect(re)
        if 'extracted_text' not in getattr(insp, "unloaded", set()):
            et = (re.extracted_text or "").strip()
            if et in ("Error extracting text.", "No readable text found."):
                return True
    except Exception:
        # Best-effort: if introspection fails, keep previous behavior as a fallback.
        et = (re.extracted_text or "").strip()
        if et in ("Error extracting text.", "No readable text found."):
            return True
    return False


def build_application_summary_response(application: Application, current_user_id: Optional[int] = None) -> ApplicationSummaryResponse:
    """Ultra-high performance summary builder using model_construct and manual mapping."""
    # 1. Manual mapping for nested models to ensure ZERO validation cost
    job_summary = None
    if application.job:
        job_summary = JobSummary.model_construct(
            id=application.job.id,
            job_id=application.job.job_id,
            title=application.job.title
        )

    re_summary = None
    if application.resume_extraction:
        re_summary = ResumeExtractionSummary.model_construct(
            id=application.resume_extraction.id,
            resume_score=application.resume_extraction.resume_score or 0.0,
            skill_match_percentage=application.resume_extraction.skill_match_percentage or 0.0,
            experience_level=application.resume_extraction.experience_level,
            summary=application.resume_extraction.summary,
            extracted_skills=application.resume_extraction.extracted_skills
        )

    int_summary = None
    if application.interview:
        # Build report summary for score bars if the report was eagerly loaded
        report_summary = None
        if application.interview.report:
            rpt = application.interview.report
            report_summary = InterviewReportSummary.model_construct(
                aptitude_score=rpt.aptitude_score,
                technical_skills_score=rpt.technical_skills_score,
                behavioral_score=rpt.behavioral_score,
            )
        int_summary = InterviewSummary.model_construct(
            id=application.interview.id,
            test_id=application.interview.test_id,
            status=application.interview.status,
            overall_score=application.interview.overall_score or 0.0,
            report=report_summary,
        )

    # 2. Ownership Calculation (Simplified per Senior Backend Engineer request)
    is_owner = (application.hr_id == current_user_id) if current_user_id else False

    # Generate full photo URL from path — use public URL (no network call) for list view.
    # Signed URLs (which require a Supabase round-trip) are only needed on the detail page.
    photo_url = None
    if application.candidate_photo_path:
        photo_url = get_public_url(settings.supabase_bucket_id_photos, application.candidate_photo_path)

    # 3. Final Construction (model_construct bypasses all Pydantic validators)
    return ApplicationSummaryResponse.model_construct(
        id=application.id,
        job_id=application.job_id,
        job=job_summary,
        candidate_name=application.candidate_name,
        candidate_email=application.candidate_email,
        candidate_phone=application.candidate_phone,
        status=application.status,
        file_status=application.file_status,
        applied_at=application.applied_at or get_ist_now(),
        resume_extraction=re_summary,
        interview=int_summary,
        resume_score=application.resume_score or 0.0,
        composite_score=application.composite_score or 0.0,
        is_owner=is_owner,
        assigned_hr_id=application.hr_id,
        assigned_hr_name=application.hr.full_name if application.hr else None,
        candidate_photo_path=application.candidate_photo_path,
        photo_url=photo_url
    )

def build_application_detail_response(application: Application, current_user_id: Optional[int] = None) -> ApplicationDetailResponse:
    from app.core.storage import get_signed_url, get_public_url
    
    resume_url = None
    if application.resume_file_path:
        resume_url = get_signed_url(settings.supabase_bucket_resumes, application.resume_file_path, expires_in=900)
    
    photo_url = None
    if application.candidate_photo_path:
        photo_url = get_signed_url(settings.supabase_bucket_id_photos, application.candidate_photo_path, expires_in=900)
             
    id_card_url = None
    if getattr(application, 'id_card_url', None):
        id_card_url = get_signed_url(settings.supabase_bucket_id_cards, application.id_card_url, expires_in=900)

    video_url = None
    if application.interview and application.interview.video_recording_path:
        video_url = get_signed_url(settings.supabase_bucket_videos, application.interview.video_recording_path, expires_in=900)

    detail = ApplicationDetailResponse.model_validate(application, from_attributes=True)
    raw_notes = application.hr_notes or ""
    degraded = RIMS_EXTRACTION_DEGRADED_MARKER in raw_notes or _heuristic_extraction_degraded(application)
    
    # Safety: Ensure file_status reflects reality (Step 5)
    inferred_file_status = application.file_status or 'active'
    if application.resume_file_path and not resume_url:
        inferred_file_status = 'missing'
    
    detail.file_status = inferred_file_status
    detail.resume_url = resume_url
    detail.photo_url = photo_url
    detail.id_card_url = id_card_url
    detail.video_url = video_url
    detail.extraction_degraded = degraded

    # Ownership Calculation (Simplified)
    if current_user_id:
        detail.is_owner = (application.hr_id == current_user_id)
    detail.assigned_hr_id = application.hr_id
    detail.assigned_hr_name = application.hr.full_name if application.hr else None

    # 7. Fallback Safety (Scenario 5 Guard)
    if not detail.applied_at:
        fallback_date = application.created_at or get_ist_now()
        detail.applied_at = fallback_date
        logger.error(
            f"INTEGRITY_ALERT: App {application.id} missing applied_at. Using fallback.",
            extra={"app_id": application.id, "fallback_date": str(fallback_date)}
        )
    
    return detail

import time
from pathlib import Path

pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

from app.core.config import get_settings
from app.core.idempotency import is_duplicate_request
from app.core.observability import get_request_id, log_json, safe_hash
settings = get_settings()

from app.core.rate_limiter import limiter
from fastapi import Request

router = APIRouter(prefix="/api/applications", tags=["applications"])


@router.get("/failures", response_model=list[ApplicationResponse])
def get_application_failures(
    db: Session = Depends(get_db),
    current_admin: User = Depends(get_current_admin)
):
    """(Control-Level) Get all applications that have failed processing."""
    from app.services.state_machine import CandidateState
    return db.query(Application).filter(
        or_(
            Application.retry_count > 0,
            Application.status == CandidateState.PERMANENT_FAILURE.value
        )
    ).order_by(Application.last_attempt_at.desc()).limit(100).all()

@router.get("/ranking/{job_id}")
def get_candidate_ranking(
    job_id: int,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Get ranked candidates for a specific job (Point 3)"""
    job = db.query(Job).filter(Job.id == job_id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    # Apply visibility isolation
    if current_user.role.lower() == "hr":
        if job.hr_id != current_user.id:
            raise HTTPException(status_code=403, detail="Forbidden: You do not own this job.")
    
    validate_hr_ownership(job, current_user, resource_name="job")

    from app.services.candidate_service import CandidateService
    service = CandidateService(db)
    ranked = service.get_ranked_candidates(job_id)
    
    result = []
    for idx, app in enumerate(ranked):
        result.append({
            "rank": idx + 1,
            "id": app.id,
            "candidate_name": app.candidate_name,
            "composite_score": app.composite_score,
            "recommendation": app.recommendation,
            "status": app.status
        })
    return result


@router.post("/apply", response_model=ApplicationResponse, status_code=status.HTTP_201_CREATED)
@limiter.limit("10/minute")  # Reduced from 300/minute to prevent public abuse. Batch uploads should use an authenticated endpoint.
async def apply_for_job(
    request: Request,
    job_id: int = Form(...),
    candidate_name: str = Form(...),
    candidate_email: Optional[str] = Form(None),
    candidate_phone: Optional[str] = Form(None),
    resume_file: UploadFile = File(...),
    photo_file: Optional[UploadFile] = File(None),
    background_tasks: BackgroundTasks = BackgroundTasks(),
    db: Session = Depends(get_db)
):
    """Apply for a job with resume (Public endpoint)"""
    # Name validation removed as requested by user.
    request_id = None
    try:
        from app.core.observability import log_json
    except ImportError:
        log_json = None
    ip_address = None
    try:
        request_id = get_request_id(request)
        if request and request.client:
            ip_address = request.client.host
    except Exception:
        request_id = None
        ip_address = None

    # 1. Email/Phone Presence Check (Point 1)
    if not candidate_email and not candidate_phone:
        raise HTTPException(
            status_code=400,
            detail="At least one valid contact method (email or phone) is required."
        )

    # 2. Email Validation & Normalization (Point 2)
    if candidate_email:
        try:
            from app.core.email_utils import validate_email_strict_enterprise

            candidate_email = validate_email_strict_enterprise(
                candidate_email,
                ip=ip_address,
                request_id=request_id,
                logger=logger,
            )
        except ValueError as e:
            logger.warning(f"Email validation failed during application creation: {e}")
            err_msg = str(e)
            if any(keyword in err_msg.lower() for keyword in ["email", "disposable", "domain"]):
                raise HTTPException(status_code=400, detail=err_msg)
            raise HTTPException(status_code=400, detail="Invalid email format.")
    else:
        candidate_email = None

    # Non-blocking flag for obviously fake/test domains (H-domain hygiene)
    suspicious_email_domain = None
    is_disposable = False # Initialize to avoid UnboundLocalError
    if candidate_email:
        try:
            from app.core.config import get_settings
            settings = get_settings()
            if settings.disposable_email_domains:
                DISPOSABLE_DOMAINS = {d.strip().lower() for d in settings.disposable_email_domains.split(",") if d.strip()}
            else:
                DISPOSABLE_DOMAINS = {
                    "fengnu.com", "mailinator.com", "guerrillamail.com", "10minutemail.com", 
                    "tempmail.com", "yopmail.com", "sharklasers.com", "getnada.com", 
                    "dispostable.com", "trashmail.com", "mail.tm", "mail.gw", "temp-mail.org"
                }
            _, domain_part = candidate_email.rsplit("@", 1)
            domain_part = domain_part.lower().strip()
            if domain_part in DISPOSABLE_DOMAINS:
                suspicious_email_domain = domain_part
                is_disposable = True
            else:
                is_disposable = False
        except Exception:
            suspicious_email_domain = None
            is_disposable = False

    # 3. Phone Validation & Normalization (Points 3, 4)
    from app.core.phone_utils import compute_phone_hash, normalize_phone_digits

    candidate_phone_raw = candidate_phone if candidate_phone else None
    candidate_phone_normalized = None
    phone_error_reason = None
    if candidate_phone:
        candidate_phone_normalized, phone_error_reason = normalize_phone_digits(candidate_phone)

        if candidate_phone_normalized is None and phone_error_reason is not None:
            try:
                from app.core.observability import log_json

                log_json(
                    logger,
                    "phone_validation_rejected",
                    request_id=request_id,
                    endpoint="/api/applications/apply",
                    user_id=None,
                    status=400,
                    level="warning",
                    extra={"reason": phone_error_reason},
                )
            except Exception:
                pass
            if phone_error_reason in ("letters_present", "invalid_length"):
                raise HTTPException(status_code=400, detail="Phone number must be 10–15 digits")
            if phone_error_reason == "invalid_characters":
                raise HTTPException(status_code=400, detail="Phone number must contain only digits (and separators)")
            raise HTTPException(status_code=400, detail="Phone number must be 10–15 digits")

    # Use correctly normalized digits for both storage and hash (Point 1)
    candidate_phone_hash = compute_phone_hash(candidate_phone_normalized) if candidate_phone_normalized else None

    # Check if job exists and is open
    job = db.query(Job).filter(Job.id == job_id, Job.status == "open").first()
    if not job:
        if log_json:
            log_json(
                logger,
                "validation_failed",
                request_id=request_id,
                endpoint="/api/applications/apply",
                status=404,
                level="warning",
                extra={"module": "applications", "field": "job_id", "reason": "not_found_or_closed", "input_preview": str(job_id)},
            )
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Job not found or not open"
        )

    request_id_header = request.headers.get("X-Request-ID")
    idempotency_key = f"{candidate_email or candidate_phone_hash}:{job_id}"
    if settings.enable_request_id_idempotency and is_duplicate_request(
        request_id=request_id_header,
        scope="applications.apply",
        key=idempotency_key.lower().strip(),
        ttl_seconds=60,
    ):
        existing_idem = (
            db.query(Application)
            .filter(
                Application.job_id == job_id,
                or_(
                    (Application.candidate_email.ilike(candidate_email)) if candidate_email else False,
                    (Application.candidate_phone_hash == candidate_phone_hash) if candidate_phone_hash else False
                )
            )
            .first()
        )
        if existing_idem:
            log_json(
                logger,
                "apply_idempotent_replay",
                request_id=request_id,
                endpoint="/api/applications/apply",
                level="info",
                extra={
                    "application_id": existing_idem.id,
                    "job_id": job_id,
                    "email_hash": safe_hash(candidate_email),
                },
            )
            return existing_idem
        raise HTTPException(
            status_code=409,
            detail="Duplicate application request detected. Please wait and retry.",
        )

    # 4. Duplicate Identification (Point 7)
    # Block any user who has already applied to this specific job using THIS email OR THIS phone.
    existing_app = db.query(Application).filter(
        Application.job_id == job_id,
        or_(
            (Application.candidate_email.ilike(candidate_email)) if candidate_email else False,
            (Application.candidate_phone_hash == candidate_phone_hash) if candidate_phone_hash else False,
        ),
    ).first()

    if existing_app:
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="You have already applied for this job.",
        )
    
    # 4. Resume Validation
    MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
    ALLOWED_RESUME_EXTENSIONS = {".pdf", ".docx", ".doc"}
    
    from app.core.resume_upload_utils import (
        generate_hashed_resume_filename,
        get_resume_extension,
        validate_resume_signature,
    )

    resume_ext = get_resume_extension(resume_file.filename)
    if resume_ext not in ALLOWED_RESUME_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Invalid resume file type. Only .pdf, .docx, and .doc are allowed.",
        )
    content = await resume_file.read()
    if not content:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid or empty file")
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File too large. Maximum size is 5MB.",
        )
    
    ok, reason = validate_resume_signature(resume_ext, content)
    if not ok:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=f"Invalid resume content: {reason}")

    # 4c. Content Validity Check (BA_007) - Detect "empty" PDFs (no readable text)
    if resume_ext == ".pdf":
        try:
            from pypdf import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            has_text = False
            for page in reader.pages:
                text_content = page.extract_text() or ""
                if text_content.strip():
                    has_text = True
                    break
            if not has_text:
                logger.warning(f"Empty PDF uploaded: {resume_file.filename}")
                raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid or empty file: No readable text found in PDF.")
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"PDF validation failed: {e}")
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid or empty file: Corrupted or unreadable PDF.")
    elif resume_ext in [".docx", ".doc"]:
        try:
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(content))
            has_text = any(para.text.strip() for para in doc.paragraphs)
            if not has_text:
                logger.warning(f"Empty DOCX uploaded: {resume_file.filename}")
                raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid or empty file: No readable text found in document.")
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"DOCX validation failed: {e}")
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Invalid or empty file: Corrupted or unreadable document.")

    # 4b. Duplicate Resume Content Detection (Point 7)
    import hashlib
    resume_hash = hashlib.sha256(content).hexdigest()
    
    existing_resume = db.query(Application).filter(
        Application.job_id == job_id,
        Application.resume_hash == resume_hash
    ).first()
    
    if existing_resume:
        logger.warning(f"Duplicate resume detected for job {job_id}. Hash: {resume_hash}. Existing app ID: {existing_resume.id}")
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="Duplicate resume entry: This file has already been processed for this job."
        )

    # 5. Photo Validation
    photo_content = None
    if photo_file:
        photo_content = await photo_file.read()
        if not photo_content:
            raise HTTPException(status_code=400, detail="Candidate photo is empty.")
        if len(photo_content) > MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail="Photo too large. Maximum size is 5MB.")
            
        # Validate JPEG/PNG magic bytes
        is_jpeg = photo_content.startswith(b"\xff\xd8\xff")
        is_png = photo_content.startswith(b"\x89PNG\r\n\x1a\n")
        if not (is_jpeg or is_png):
            raise HTTPException(
                status_code=400,
                detail="Invalid photo format. Only JPEG and PNG images are allowed."
            )

    # 6. Database Entry (Atomic Transaction)
    # We flush first to get the ID, then upload to storage using that ID as a prefix.
    user_agent = request.headers.get("user-agent")
    filename = generate_hashed_resume_filename(
        candidate_email=candidate_email or f"phone_{candidate_phone_hash}",
        job_id=job_id,
        resume_ext=resume_ext,
        content=content,
    )

    resume_storage_path = None
    photo_storage_path = None

    warning_notes = None
    if suspicious_email_domain:
        warning_notes = f"Warning: Possibly fake/test email domain detected ({suspicious_email_domain})."

    new_application = Application(
        job_id=job_id,
        hr_id=job.hr_id,
        candidate_name=candidate_name,
        candidate_email=candidate_email,
        candidate_phone=candidate_phone_normalized,
        candidate_phone_normalized=None,
        candidate_phone_raw=None,
        candidate_phone_hash=candidate_phone_hash,
        resume_file_name=resume_file.filename,
        resume_hash=resume_hash,
        status="applied",
        hr_notes=warning_notes,
        is_disposable_email=is_disposable,
        applied_at=datetime.now(timezone(timedelta(hours=5, minutes=30))).replace(tzinfo=None),
        resume_status="pending",
    )

    try:
        db.add(new_application)
        db.flush() # Get ID for storage paths

        # Upload Resume to Supabase
        resume_storage_path = f"{new_application.id}/resume_{filename}"
        returned_resume_path = upload_file(settings.supabase_bucket_resumes, resume_storage_path, content, content_type=resume_file.content_type)
        if not returned_resume_path:
            raise Exception("Resume storage upload failed")
        new_application.resume_file_path = returned_resume_path

        # Upload Photo to Supabase
        if photo_content:
            # Derive extension from magic bytes, not filename
            photo_ext = ".jpg" # default
            if photo_content.startswith(b'\xff\xd8\xff'):
                photo_ext = ".jpg"
            elif photo_content.startswith(b'\x89PNG'):
                photo_ext = ".png"
            elif photo_content.startswith(b'GIF87a') or photo_content.startswith(b'GIF89a'):
                photo_ext = ".gif"
            elif photo_content.startswith(b'RIFF') and photo_content[8:12] == b'WEBP':
                photo_ext = ".webp"
            
            photo_storage_path = f"{new_application.id}/photo_initial{photo_ext}"
            returned_photo_path = upload_file(settings.supabase_bucket_id_photos, photo_storage_path, photo_content, content_type=photo_file.content_type)
            if not returned_photo_path:
                raise Exception("Candidate photo storage upload failed")
            new_application.candidate_photo_path = returned_photo_path

        # Create HR Notification
        from app.domain.models import Notification
        db.add(Notification(
            user_id=job.hr_id,
            notification_type="NEW_APPLICATION",
            title=f"New Application: {candidate_name}",
            message=f"{candidate_name} has applied for {job.title}.",
            related_application_id=new_application.id,
        ))

        db.commit()
        db.refresh(new_application)

        # Audit Log
        from app.services.candidate_service import CandidateService
        CandidateService(db).create_audit_log(None, "APPLICATION_CREATED", "Application", new_application.id, {"job_id": job_id})

        # [Trigger] Application Received Email
        try:
            if not new_application or not new_application.candidate_email:
                logger.error(f"[EMAIL][FAILED] Missing email for App #{getattr(new_application, 'id', 'UNKNOWN')}")
            elif getattr(new_application, "_email_sent", False):
                logger.warning(f"[EMAIL][SKIPPED] Duplicate prevented for App #{new_application.id}")
            else:
                background_tasks.add_task(send_application_received_email, new_application)
                new_application._email_sent = True
                logger.info(f"[EMAIL] Application email queued for App #{new_application.id}")
        except Exception as e:
            logger.error(f"[EMAIL][FAILED] Application email for App #{getattr(new_application, 'id', 'UNKNOWN')}: {str(e)}")


    except IntegrityError as e:
        db.rollback()
        logger.warning(f"Duplicate application integrity violation: {e}")
        from app.core.storage import delete_file
        try:
            if resume_storage_path:
                delete_file(settings.supabase_bucket_resumes, resume_storage_path)
            if photo_storage_path:
                delete_file(settings.supabase_bucket_id_photos, photo_storage_path)
        except Exception as del_err:
            logger.warning(f"Failed to clean up files on IntegrityError: {del_err}")
        raise HTTPException(
            status_code=status.HTTP_409_CONFLICT,
            detail="You have already applied for this job."
        )
    except Exception as e:
        db.rollback()
        logger.error(f"Application submission failed: {e}")
        # Clean up cloud files if they were uploaded
        from app.core.storage import delete_file
        try:
            if resume_storage_path and new_application.resume_file_path:
                delete_file(settings.supabase_bucket_resumes, resume_storage_path)
            if photo_storage_path and new_application.candidate_photo_path:
                delete_file(settings.supabase_bucket_id_photos, photo_storage_path)
        except Exception as cleanup_err:
            logger.warning(f"Failed to clean up files for application: {cleanup_err}")
        raise HTTPException(status_code=500, detail="Failed to submit application securely.")

    background_tasks.add_task(
        process_application_background, 
        new_application.id, 
        job_id, 
        new_application.resume_file_path, 
        candidate_email, 
        candidate_name
    )
    
    return new_application

async def process_application_background(application_id: int, job_id: int, abs_file_path: str, candidate_email: str, candidate_name: str):
    """Heavy AI processing and notification workflow in background.
    Serialized through _ai_analysis_semaphore to prevent concurrent Groq API calls
    during batch uploads (which would cause rate-limit / 503 failures).
    """
    db = SessionLocal()
    try:
        from app.services.candidate_service import CandidateService
        cand_service = CandidateService(db)
        
        # Reload objects in this session
        # Step 1: lock only (without joins)
        # Note: Application.resume_extraction is configured as lazy='joined', so ORM-level
        # .with_for_update() may turn into a LEFT OUTER JOIN, which Postgres rejects.
        db.execute(text("SELECT 1 FROM applications WHERE id = :id FOR UPDATE"), {"id": application_id})
        
        # Step 2: fetch with joins, no lock
        application = db.query(Application).options(
            joinedload(Application.resume_extraction)
        ).filter(Application.id == application_id).first()
        
        job = db.query(Job).filter(Job.id == job_id).first()
        if not application or not job:
            db.close()
            return

        application.resume_status = "parsing"
        application.parsing_started_at = get_ist_now()
        db.commit()
        db.refresh(application)
        try:
            log_json(
                logger,
                "resume_parsing_started",
                level="info",
                extra={"application_id": application_id, "job_id": job_id},
            )
        except Exception:
            pass

        # 1. Initial State
        cand_service.advance_stage(application_id, "Application Submitted", "pass")
        cand_service.create_audit_log(None, "APPLICATION_SUBMITTED", "Application", application_id, {"email": candidate_email})
        
        # 2. Screening Stage
        cand_service.advance_stage(application_id, "Resume Screening", "pending")
        
        # Parse resume text based on file type from Supabase Storage
        resume_text = ""
        try:
            if not abs_file_path:
                raise Exception("No resume file path available (storage likely disabled)")
                
            from app.core.storage import get_supabase_client
            supabase = get_supabase_client()
            if not supabase:
                raise Exception("Supabase client not initialized")
            
            # Download from storage
            # Assuming 'resumes' bucket
            bucket_name = settings.supabase_bucket_resumes
            from app.core.storage import download_file
            response = download_file(bucket_name, abs_file_path)
            if not response:
                raise Exception(f"Failed to download resume from {abs_file_path}")
            
            from io import BytesIO
            file_stream = BytesIO(response)
            
            file_ext = abs_file_path.lower().split('_')[-1].split('.')[-1] if '.' in abs_file_path else 'pdf'
            
            if file_ext == 'pdf':
                # Handle potential BOM or leading whitespace in binary response
                if response.lstrip().startswith(b'%PDF'):
                    from pypdf import PdfReader
                    reader = PdfReader(file_stream)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            resume_text += page_text + "\n"
                else:
                    # Fallback for plain-text or disguised PDFs
                    resume_text = response.decode('utf-8', errors='ignore')
            elif file_ext == 'docx':
                import docx
                doc = docx.Document(file_stream)
                for para in doc.paragraphs:
                    if para.text:
                        resume_text += para.text + "\n"
            elif file_ext == 'doc':
                # python-docx doesn't support legacy .doc; fallback to raw text if it looks readable
                resume_text = response.decode('utf-8', errors='ignore')
            else:
                resume_text = response.decode('utf-8', errors='ignore')

            # Post-extraction sanity check for scanned PDFs
            if file_ext == 'pdf' and len(response) > 50000 and len(resume_text.strip()) < 100:
                resume_text = "[[SCANNED_PDF_DETECTED]]\n" + resume_text
        except Exception as e:
            logger.error(f"Background Text Extraction Skipped or Failed: {e}")
            cand_service.create_audit_log(None, "RESUME_TEXT_EXTRACTION_SKIPPED", "Application", application_id, {"reason": str(e)})
            resume_text = "Parsing skipped: Storage unavailable or file missing."
        
        if not resume_text.strip() or resume_text.startswith("Parsing skipped:"):
            # Skip AI if no text or explicitly skipped due to streaming/parsing failure
            extraction_data = {
                "summary": "AI extraction skipped: " + resume_text,
                "skills": ["Unparsable File"],
                "experience": 0,
                "score": 0.0,
                "match_percentage": 0,
                "extraction_degraded": True,
                "is_resume": False,
                "reasoning": {"ai_justification": "System detected a corrupted, missing, or unparsable file. Standard AI evaluation was skipped."}
            }
            extraction_degraded_flag = True
        else:
            # AI Parsing — serialized through semaphore to avoid concurrent Groq rate-limit errors
            async with _ai_analysis_semaphore:
                extraction_data = await parse_resume_with_ai(resume_text, job_id, job.description, job.experience_level)
            extraction_degraded_flag = extraction_data.pop("extraction_degraded", False)

        # Store extraction (Versioning + Upsert pattern)
        resume_extraction = db.query(ResumeExtraction).filter(ResumeExtraction.application_id == application_id).first()
        if resume_extraction:
            # ── Phase 3: Versioning (Save old record before overwrite) ──
            try:
                version_count = db.query(ResumeExtractionVersion).filter(ResumeExtractionVersion.application_id == application_id).count()
                old_version = ResumeExtractionVersion(
                    application_id=application_id,
                    version_number=version_count + 1,
                    extracted_text=resume_extraction.extracted_text,
                    extracted_skills=resume_extraction.extracted_skills,
                    resume_score=resume_extraction.resume_score
                )
                db.add(old_version)
            except Exception as e:
                logger.warning(f"Failed to version old resume extraction: {e}")
        else:
            resume_extraction = ResumeExtraction(application_id=application_id)
            db.add(resume_extraction)
            
        resume_extraction.extracted_text = resume_text
        resume_extraction.summary = extraction_data.get("summary", "")
        resume_extraction.extracted_skills = json.dumps(extraction_data.get("skills") or [])
        resume_extraction.years_of_experience = extraction_data.get("experience")
        resume_extraction.education = json.dumps(extraction_data.get("education") or [])
        resume_extraction.previous_roles = json.dumps(extraction_data.get("roles") or [])
        resume_extraction.experience_level = extraction_data.get("experience_level")
        resume_extraction.resume_score = extraction_data.get("score", 0)
        resume_extraction.skill_match_percentage = extraction_data.get("match_percentage", 0)
        resume_extraction.reasoning = {"ai_justification": extraction_data.get("reasoning")}
        
        if extraction_data.get("candidate_name"):
            resume_extraction.candidate_name = extraction_data.get("candidate_name")
        if extraction_data.get("email"):
            resume_extraction.email = extraction_data.get("email")
        if extraction_data.get("phone_number"):
            resume_extraction.phone_number = extraction_data.get("phone_number")
        
        # Update Application summary fields
        application.resume_score = extraction_data.get("score", 0)
        
        # ── Phase 7: Scoring Transparency ──
        application.scoring_metadata = {
            "logic_version": "v2.0",
            "weights": {"skills": 0.6, "experience": 0.4},
            "recomputed_at": get_ist_now().isoformat(),
            "extraction_degraded": extraction_degraded_flag
        }
        
        # ── HYBRID IDENTITY EXTRACTION: AI + regex fallback ──
        from app.services.ai_service import extract_email_regex, extract_phone_regex, extract_name_heuristic
        
        extracted_name = extraction_data.get("candidate_name") or extract_name_heuristic(resume_text)
        extracted_email = extraction_data.get("email") or extract_email_regex(resume_text)
        extracted_phone = extraction_data.get("phone_number") or extract_phone_regex(resume_text)
        
        logger.info(f"[IDENTITY SYNC] App #{application_id} | EXTRACTED: name={extracted_name}, email={extracted_email}, phone={extracted_phone}")
        
        is_duplicate = False
        duplicate_app_id = None
        
        # Check for placeholder status BEFORE any updates
        is_emailed_app = application.hr_notes == "Ingested automatically from Email Recruiter Channel."
        email_is_placeholder = (application.candidate_email and "@batch." in application.candidate_email) or is_emailed_app
        name_is_placeholder = not application.candidate_name or len(application.candidate_name.split()) < 2 or is_emailed_app
        
        # 1. Duplicate Detection via Extracted Email (Point 1: Prevent Clashes)
        if extracted_email:
            norm_email = extracted_email.lower().strip()
            # Heuristic: Check if this email exists anywhere in the system (global check)
            # OR if it was already extracted for this job in a concurrent batch task.
            existing_match = db.query(Application).outerjoin(ResumeExtraction).filter(
                Application.job_id == application.job_id,
                or_(
                    Application.candidate_email == norm_email,
                    ResumeExtraction.email == norm_email
                ),
                Application.id != application_id
            ).first()
            
            if existing_match:
                logger.info(f"[IDENTITY SYNC] App #{application_id} Conflict: '{norm_email}' already exists in system. Deleting duplicate.")
                
                # Clean up cloud files before deleting DB record
                from app.core.storage import delete_file
                try:
                    if application.resume_file_path:
                        delete_file(settings.supabase_bucket_resumes, application.resume_file_path)
                    if application.candidate_photo_path:
                        delete_file(settings.supabase_bucket_id_photos, application.candidate_photo_path)
                except Exception as e:
                    logger.warning(f"Failed to clean up files for duplicate application #{application_id}: {e}")

                # Delete related resume_extraction first if it exists in session to avoid ForeignKeyViolation
                if resume_extraction:
                    try:
                        insp = sa_inspect(resume_extraction)
                        if insp.persistent:
                            db.delete(resume_extraction)
                        elif resume_extraction in db:
                            db.expunge(resume_extraction)
                    except Exception:
                        pass
                
                db.delete(application)
                db.commit()
                db.close()
                return
            else:
                # Safe to update email if it was missing or obviously wrong (placeholder)
                if not application.candidate_email or email_is_placeholder:
                    application.candidate_email = extracted_email
                    logger.info(f"[IDENTITY SYNC] App #{application_id} | Email updated from resume")

        # 2. Name & Phone Sync (Conservative)
        # We only overwrite if the current value is "weak" (one word name or placeholder from batch)
        if extracted_name and (name_is_placeholder or email_is_placeholder):
            application.candidate_name = extracted_name
            logger.info(f"[IDENTITY SYNC] App #{application_id} | Name updated from resume")
            
        if extracted_phone and not application.candidate_phone:
            application.candidate_phone = extracted_phone
            logger.info(f"[IDENTITY SYNC] App #{application_id} | Phone updated from resume")
            
        if extracted_phone and not application.candidate_phone:
            application.candidate_phone = extracted_phone
            logger.info(f"[IDENTITY SYNC] App #{application_id} | Phone updated from resume")

        if extraction_degraded_flag:
            _append_extraction_degraded_marker(application)

        application.resume_status = "parsed"; application.failure_reason = None
        db.commit()
        db.refresh(application)
        try:
            log_json(
                logger,
                "resume_parsing_completed",
                level="info",
                extra={"application_id": application_id, "resume_status": "parsed"},
            )
        except Exception:
            pass

        # ── Pipeline Advancement ──
        # If duplicate, we mark stage as 'fail' so it's not advanced, but we STILL provide the real score.
        # This is *internal-only* and never surfaces as a 409 or candidate-facing error.
        # If not duplicate, we mark as 'pass' to await HR decision.
        stage_status = "fail" if is_duplicate else "pass"
        stage_note = (
            f"Possible duplicate of App #{duplicate_app_id}"
            if is_duplicate
            else "AI analysis complete — awaiting HR decision"
        )
        
        cand_service.advance_stage(
            application_id, 
            "Resume Screening", 
            stage_status, 
            extraction_data.get("score", 0) * 10, 
            stage_note
        )
        
        # ── Phase 6: Critical Audit Logging ──
        cand_service.create_audit_log(
            None, 
            "RESUME_SCREENING_COMPLETED", 
            "Application", 
            application_id, 
            {"score": extraction_data.get("score", 0), "match": extraction_data.get("match_percentage", 0)},
            is_critical=True
        )
        
        db.commit()
    except Exception as e:
        logger.error(
            f"CRITICAL Background Error processing application {application_id}: {e}",
            exc_info=True,
        )
        db.rollback()
        # Log the critical error
        try:
            cand_service = CandidateService(db) # Re-initialize if needed, or pass db
            cand_service.create_audit_log(None, "BACKGROUND_PROCESSING_FAILED", "Application", application_id, {"error": str(e)})
            # Optionally update application status to indicate processing failed
            application = db.query(Application).filter(Application.id == application_id).with_for_update().first()
            if application:
                # Only reset to 'applied' if it hasn't been advanced by HR (Fix for status-sync bug)
                if application.status == "applied":
                    application.status = "applied" 
                application.resume_status = "failed"
                application.retry_count = (application.retry_count or 0) + 1
                application.failure_reason = str(e)[:1000]
                application.last_attempt_at = get_ist_now()

                # Escalation (Phase 5 fix)
                if application.retry_count >= 3:
                     from app.services.state_machine import CandidateState
                     application.status = CandidateState.PERMANENT_FAILURE.value
                     application.failure_reason = "[PERMANENT_FAILURE]: " + application.failure_reason

                # Error details are now handled by the frontend via failure_reason
                _append_extraction_degraded_marker(application)
                db.commit()
                try:
                    log_json(
                        logger,
                        "resume_parsing_failed",
                        level="error",
                        extra={"application_id": application_id, "resume_status": "failed"},
                    )
                except Exception:
                    pass
        except Exception as log_e:
            logger.error(f"Failed to log critical error for application {application_id}: {log_e}")
    finally:
        db.close()
    
    # The return value of a background task is not used by FastAPI.
    # The original snippet returned new_application, but it's not necessary here.
    # Keeping it for consistency with the provided snippet, but it has no effect.
    return application


@router.get("/has-applied", response_model=HasAppliedResponse)
@limiter.limit("5/minute")  # Stricter rate limit to prevent email/phone enumeration
def has_applied_for_job(
    request: Request,
    job_id: int,
    candidate_email: str,
    candidate_phone: str, # Require both email and phone to confirm application status
    db: Session = Depends(get_db),
):
    """Return whether a (job_id, candidate_email/phone) application already exists.
    
    CRITICAL: Validates against both email and normalized phone hash to strictly
    enforce one application per person per job.
    """
    from app.core.email_utils import validate_email_strict_enterprise
    from app.core.phone_utils import compute_phone_hash, normalize_phone_digits

    # 1. Normalize Email
    if candidate_email:
        try:
            candidate_email = validate_email_strict_enterprise(
                candidate_email,
                ip=None,
                request_id=None,
                logger=logger,
            )
        except ValueError:
            # If email is invalid, it can't match a stored valid email
            candidate_email = None

    # 2. Normalize and Hash Phone
    candidate_phone_hash = None
    if candidate_phone:
        normalized_digits, _ = normalize_phone_digits(candidate_phone)
        if normalized_digits:
            candidate_phone_hash = compute_phone_hash(normalized_digits)

    # 3. Check for duplicates (OR logic)
    if not candidate_email and not candidate_phone_hash:
        db.query(Application.id).filter(Application.id == -1).first() # Dummy query to equalize timing
        return HasAppliedResponse(hasApplied=False)

    existing = (
        db.query(Application.id)
        .filter(
            Application.job_id == job_id,
            or_(
                (Application.candidate_email.ilike(candidate_email)) if candidate_email else False,
                (Application.candidate_phone_hash == candidate_phone_hash) if candidate_phone_hash else False,
            ),
        )
        .first()
    )
    return HasAppliedResponse(hasApplied=existing is not None)


@router.get("/pending-count")
def get_pending_applications_count(
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db),
):
    """Sidebar / badges: count applications not in a terminal state, scoped to HR's jobs."""
    # Use func.count() directly — avoids loading full ORM objects just to count them.
    q = db.query(func.count(Application.id)).filter(
        ~Application.status.in_(("onboarded", "rejected", "permanent_failure")),
        or_(func.trim(Application.file_status).in_(('active', 'missing')), Application.file_status == None)
    )

    # Apply visibility isolation: Anyone not a super_admin is restricted to their own data
    if current_user.role.lower() not in ["super_admin", "admin"]:
        q = q.outerjoin(Job, Application.job_id == Job.id)
        q = q.filter(or_(Job.hr_id == current_user.id, Application.hr_id == current_user.id))

    count = q.scalar() or 0
    return {"count": count}


@router.get("", response_model=ApplicationListResponse, response_class=ORJSONResponse)
def get_hr_applications(
    job_id: str = None,
    from_date: str = None,
    to_date: str = None,
    status: str = None,
    time_range: str = None,
    search: str = None,
    skip: int = 0,
    limit: int = Query(default=50, le=1000),
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Get all applications for HR's jobs (HR only)"""
    t_start = time.perf_counter()
    items = []
    total = 0
    safe_skip = max(0, int(skip or 0))
    safe_limit = max(1, min(int(limit or 1000), 1000))
    
    try:
        # 1. Build a base query with only the outerjoin and filters — no options/joinedloads.
        # This is used for the COUNT so we don't pay the cost of joinedloads twice.
        # Ensure we only include applications for jobs that have a valid job_id, or are the internal demo job
        base_query = db.query(Application).join(Job, Application.job_id == Job.id).filter(
            or_(
                and_(Job.job_id.isnot(None), Job.job_id != ""),
                Job.title == "INTERNAL_DEMO_JOB"
            )
        )

        # 2. Filters (applied to base_query so both count and data share the same filters)
        if job_id and str(job_id).strip() not in ("all", "All"):
            job_id_str = str(job_id).strip()
            if job_id_str.isdigit():
                base_query = base_query.filter(Application.job_id == int(job_id_str))
            else:
                base_query = base_query.filter(Job.job_id == job_id_str)

        if status and status != 'all':
            if status == "applied":
                base_query = base_query.filter(Application.status.in_(("applied", "submitted")))
            else:
                base_query = base_query.filter(Application.status == status)

        if search and str(search).strip():
            search_terms = str(search).strip().split()
            if search_terms:
                base_query = base_query.outerjoin(Interview, Application.id == Interview.application_id)
                for term in search_terms:
                    t = f"%{term}%"
                    base_query = base_query.filter(or_(
                        Application.candidate_name.ilike(t),
                        Application.candidate_email.ilike(t),
                        Job.title.ilike(t),
                        Job.job_id.ilike(t),
                        Interview.test_id.ilike(t)
                    ))

        def parse_date(date_str):
            if not date_str: return None
            for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
                try:
                    return datetime.strptime(date_str, fmt).date()
                except ValueError:
                    continue
            return None

        if from_date:
            sd = parse_date(from_date)
            if sd:
                logger.info(f"Applying from_date filter (IST): {sd}")
                base_query = base_query.filter(func.date(Application.applied_at) >= sd)
            else:
                logger.warning(f"Invalid from_date format: {from_date}")

        if to_date:
            ed = parse_date(to_date)
            if ed:
                logger.info(f"Applying to_date filter (IST): {ed}")
                base_query = base_query.filter(func.date(Application.applied_at) <= ed)
            else:
                logger.warning(f"Invalid to_date format: {to_date}")

        if time_range and time_range != 'all':
            # Applied Time Window Filter (Hour of Day)
            # Adjusted for IST (UTC+5:30) as project context (e.g. phone defaults) is India-based.
            # IST 06:00-12:00 (Morning)   => UTC 00:30-06:30
            # IST 12:00-18:00 (Afternoon) => UTC 06:30-12:30
            # IST 18:00-00:00 (Evening)   => UTC 12:30-18:30
            # IST 00:00-06:00 (Night)     => UTC 18:30-00:30

            # Data is now stored in IST, extract hour directly
            ist_hour = extract('hour', Application.applied_at)

            if time_range == 'morning':
                base_query = base_query.filter(ist_hour >= 6, ist_hour < 12)
            elif time_range == 'afternoon':
                base_query = base_query.filter(ist_hour >= 12, ist_hour < 18)
            elif time_range == 'evening':
                base_query = base_query.filter(ist_hour >= 18, ist_hour < 24)
            elif time_range == 'night':
                base_query = base_query.filter(ist_hour < 6)

        # 3. Security: Apply visibility isolation (HR sees only their own apps/jobs)
        if current_user.role.lower() not in ["super_admin", "admin"]:
            # Standard HR sees jobs they own OR apps they are assigned to.
            # Job is already joined via outerjoin at the start of the query.
            base_query = base_query.filter(or_(Job.hr_id == current_user.id, Application.hr_id == current_user.id))
        # Super Admin sees all.

        # 4. Count — runs on base_query (no joinedloads), which is much cheaper.
        total = base_query.count()

        # 5. Data fetch — add all eager-loading options only for the actual data query.
        #    FIX: Application.hr_id and Application.job_id MUST be in load_only; without them
        #    SQLAlchemy issues a lazy SELECT per application (N+1) for ownership checks.
        data_query = base_query.options(
            joinedload(Application.job).load_only(Job.id, Job.title, Job.hr_id, Job.status, Job.job_id),
            joinedload(Application.hr).load_only(User.id, User.full_name),
            joinedload(Application.resume_extraction).load_only(
                ResumeExtraction.id, ResumeExtraction.resume_score,
                ResumeExtraction.skill_match_percentage, ResumeExtraction.experience_level,
                ResumeExtraction.summary, ResumeExtraction.extracted_skills,
            ),
            load_only(
                Application.id, Application.job_id, Application.hr_id,
                Application.candidate_name, Application.candidate_email,
                Application.status, Application.applied_at, Application.file_status,
                Application.candidate_photo_path, Application.resume_file_path,
                Application.resume_score, Application.composite_score,
            ),
            # Load interview with test_id for badge display; also eagerly load the report
            # so that score bars (aptitude/technical/behavioral) are populated without lazy fetches.
            joinedload(Application.interview).options(
                load_only(Interview.id, Interview.test_id, Interview.status, Interview.overall_score),
                joinedload(Interview.report).load_only(
                    InterviewReport.aptitude_score,
                    InterviewReport.technical_skills_score,
                    InterviewReport.behavioral_score,
                ),
            ),
        )
        applications = data_query.order_by(Application.applied_at.desc(), Application.id.desc()).offset(safe_skip).limit(safe_limit).all()

        # Path sanitization
        for app in applications:
            for field in ['candidate_photo_path', 'resume_file_path']:
                val = getattr(app, field)
                if val and "uploads" in val:
                    idx = val.find("uploads")
                    setattr(app, field, val[idx:].replace("\\", "/"))

        # 5. Response Mapping
        t_map_start = time.perf_counter()
        items = [build_application_summary_response(app, current_user.id) for app in applications]
        t_map_duration = time.perf_counter() - t_map_start
        
        duration = time.perf_counter() - t_start
        logger.info(f"PERFORMANCE_TRACE: get_hr_applications total={duration:.4f}s (Map: {t_map_duration:.4f}s)")
        
        pages = (total + safe_limit - 1) // safe_limit
        return {
            "items": items,
            "total": total,
            "page": (safe_skip // safe_limit) + 1,
            "size": safe_limit,
            "pages": pages
        }
    except Exception as e:
        logger.error(f"APPLICATION_API_ERROR: {str(e)}", exc_info=True)
        return {"items": [], "total": 0, "page": 1, "size": safe_limit, "pages": 0, "error_hint": str(e)}

@router.get("/{application_id}/resume/download")
def download_resume(
    application_id: int,
    request: Request,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Securely download a candidate's resume (HR only)"""
    application = db.query(Application).filter(Application.id == application_id).with_for_update().first()
    if not application or not application.resume_file_path:
        raise HTTPException(status_code=404, detail="Resume file not found")
    validate_hr_ownership(application, current_user, resource_name="application")

    from fastapi.responses import RedirectResponse, FileResponse
    stored_path = (application.resume_file_path or "").replace("\\", "/")
    
    # 1. Cloud Storage Redirect (New)
    if "uploads" not in stored_path:
        from app.core.storage import get_signed_url
        signed_url = get_signed_url(settings.supabase_bucket_resumes, application.resume_file_path)
        if signed_url:
            return RedirectResponse(url=signed_url)
    
    # 2. Local File Fallback (Legacy)
    filename = os.path.basename(stored_path)
    candidate_1 = settings.uploads_dir / "resumes" / filename
    candidate_2 = None
    if "uploads/" in stored_path:
        rel = stored_path.split("uploads/", 1)[1]
        candidate_2 = settings.uploads_dir / rel
    candidate_3 = settings.uploads_dir / filename

    file_path = None
    for c in [candidate_1, candidate_2, candidate_3]:
        if c and c.exists():
            file_path = c
            break

    if not file_path:
        raise HTTPException(status_code=404, detail="Resume file not found on server")

    return FileResponse(
        path=str(file_path),
        filename=application.resume_file_name or filename,
        media_type='application/octet-stream'
    )



# --- Email Ingestion API Endpoints (Phase 2) ---
from app.services.email_ingestion_service import fetch_resume_attachments, run_batch_resume_processing
from pydantic import BaseModel

class EmailIngestRequest(BaseModel):
    # BUG-001 Fix: Credentials are no longer accepted in the request body.
    # The server reads IMAP credentials exclusively from the encrypted GlobalSettings table.
    trigger: bool = True

from app.domain.models import AttachmentResume

@router.post("/ingest-emails")
@limiter.limit("2/minute")  # BUG-004 Fix: Rate limit to prevent IMAP connection flooding
async def ingest_email_resumes(
    request: Request,
    req: EmailIngestRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """
    Trigger manual email ingestion via IMAP and map/analyze them immediately.
    """
    from app.core.encryption import decrypt_field

    imap_email = current_user.imap_email or ""
    enc_password = current_user.imap_password or ""
    
    imap_password = decrypt_field(enc_password) if enc_password else ""

    if not imap_email or not imap_password:
        raise HTTPException(
            status_code=400,
            detail="IMAP credentials not configured. Please save mailbox settings first."
        )

    # Both fetch and AI processing are moved to background to prevent Gunicorn worker timeouts
    # when processing large batches of emails with heavy attachments.
    async def run_sync_in_background():
        bg_db = SessionLocal()
        try:
            logger.info("Starting background IMAP fetch...")
            fetch_result = fetch_resume_attachments(bg_db, imap_email, imap_password, hr_id=current_user.id)
            if fetch_result and fetch_result.get("success"):
                logger.info("Background IMAP fetch complete. Starting AI mapping...")
                await run_batch_resume_processing(bg_db, hr_id=current_user.id)
            else:
                logger.error(f"Background IMAP fetch failed: {fetch_result.get('error') if fetch_result else 'Unknown error'}")
        except Exception as e:
            logger.error(f"Background email sync failed unexpectedly: {e}")
        finally:
            bg_db.close()
            
    background_tasks.add_task(run_sync_in_background)
    
    logger.info("✅ Email sync task queued in background")
    
    return {
        "success": True,
        "saved_count": 0,
        "message": "Mailbox sync started in the background. Resumes will appear here shortly.",
        "processing_triggered": True
    }

def _extract_storage_path_identifier(path: str) -> Optional[str]:
    """
    Extract a unique identifier from a storage path or URL.
    Works for both 'resumes' and 'MAIL_ATTACHMENTS' buckets.
    """
    if not path:
        return None
    # 1. Remove query parameters
    path = path.split("?")[0]
    # 2. Extract the part after the bucket name if present
    # We include the bucket prefix to avoid collisions between different buckets
    if "/MAIL_ATTACHMENTS/" in path:
        return "MAIL_ATTACHMENTS/" + path.split("/MAIL_ATTACHMENTS/")[-1].strip("/")
    if "/resumes/" in path:
        return "resumes/" + path.split("/resumes/")[-1].strip("/")
    
    # 3. Handle relative paths
    if path.startswith("MAIL_ATTACHMENTS/"):
        return path.strip("/")
    if path.startswith("resumes/"):
        return path.strip("/")
        
    return None # Don't return generic paths


def _get_target_job_id_from_subject(subject_str: str, db: Session) -> Optional[int]:
    if not subject_str:
        return None
    # Pattern A: JOB-[A-Z0-9]{6}
    job_codes = re.findall(r'JOB-[A-Z0-9]{6}', subject_str, re.IGNORECASE)
    if job_codes:
        for code in job_codes:
            extracted_code = code.upper().strip()
            job = db.query(Job).filter(func.upper(Job.job_id) == extracted_code).first()
            if job:
                return job.id
    
    # Pattern B: job id/code numeric
    subject_lower = subject_str.lower()
    numeric_id_match = re.search(r'\bjob\s*(?:id|code)?\s*[:\-\#]?\s*(\d+)\b', subject_lower)
    if numeric_id_match:
        extracted_id = int(numeric_id_match.group(1).strip())
        job = db.query(Job).filter(Job.id == extracted_id).first()
        if job:
            return job.id
            
    # Pattern C: Job title matching 80%
    subject_words = set(subject_lower.split())
    open_jobs = db.query(Job).filter(Job.status == 'open').all()
    for job in open_jobs:
        job_title_words = set(job.title.lower().split())
        if job_title_words:
            match_count = len(job_title_words & subject_words)
            if match_count / len(job_title_words) >= 0.8:
                return job.id
    return None

def _resolve_resume_mapping(
    item: AttachmentResume,
    db: Session,
    app_paths_set: set = None,
    app_emails_no_path_set: set = None,
    app_emails_job_set: set = None
):
    """
    Unified logic to resolve if an ingested email is mapped to an application.

    BUG-H Fix: Two distinct return-type contracts, now explicitly documented:
      - With sets (stats/fast path):  returns (bool, bool)  — just truthy/falsy
      - Without sets (list/slow path): returns (Application|None, bool) — full obj needed for metadata

    Callers must handle accordingly:
      - Stats call (with sets): `if app_found or is_processed`
      - Listing call (without sets): `app_obj.id if app_obj else None`
    """
    # Extract candidate's raw email from sender email
    sender_str = item.sender_email or ""
    match = re.search(r'<([^>]+)>', sender_str)
    raw_email = match.group(1).lower().strip() if match else sender_str.lower().strip()

    path_id = _extract_storage_path_identifier(item.file_url) if item.file_url else None

    # 1. Match by unique storage path identifier (most accurate)
    if path_id:
        if app_paths_set is not None:
            # Fast path (stats): return bool, bool
            if path_id in app_paths_set:
                return True, True
        else:
            # Slow path (list): return full Application object
            app = db.query(Application).filter(
                or_(
                    Application.resume_file_path == path_id,
                    Application.resume_file_path.like(f"%{path_id}%")
                )
            ).first()
            if app:
                return app, True

    # 2. Match by email (if email has no attachment and was manually mapped)
    if not path_id and raw_email:
        if app_emails_no_path_set is not None:
            if raw_email in app_emails_no_path_set:
                return True, True
        else:
            app = db.query(Application).filter(
                func.trim(func.lower(Application.candidate_email)) == raw_email,
                or_(
                    Application.resume_file_path == None,
                    Application.resume_file_path == ""
                )
            ).first()
            if app:
                return app, True

    # 3. Match by candidate email and target job ID (fallback)
    # This prevents any candidate email from mapping to their other unrelated jobs.
    target_job_id = _get_target_job_id_from_subject(item.subject, db)
    if raw_email and target_job_id:
        if app_emails_job_set is not None:
            # Fast path (stats): return bool, bool
            if (raw_email, target_job_id) in app_emails_job_set:
                return True, True
        else:
            # Slow path (list): return full Application object
            app = db.query(Application).filter(
                func.trim(func.lower(Application.candidate_email)) == raw_email,
                Application.job_id == target_job_id
            ).first()
            if app:
                return app, True

    return None, getattr(item, 'processed', False)



def _resolve_resume_mapping_bool(
    item: AttachmentResume,
    db: Session,
    app_paths_set: set,
    app_emails_no_path_set: set,
    app_emails_job_set: set
) -> bool:
    app_found, is_processed = _resolve_resume_mapping(
        item, db, app_paths_set, app_emails_no_path_set, app_emails_job_set
    )
    return bool(app_found)


def _is_email_visible_to_user(
    item: AttachmentResume,
    current_user: User,
    db: Session,
    hr_app_paths: set,
    hr_app_emails_no_path: set,
    hr_app_emails_job: set,
    global_app_paths: set,
    global_app_emails_no_path: set,
    global_app_emails_job: set,
    hr_job_ids: set
) -> tuple[bool, bool]:
    """
    Returns (is_visible, is_mapped) for the given email item.
    """
    if current_user.role in ["super_admin", "admin"]:
        is_mapped = _resolve_resume_mapping_bool(
            item, db, global_app_paths, global_app_emails_no_path, global_app_emails_job
        )
        return True, is_mapped

    # Regular HR user
    # 1. Visible if fetched by this HR manager
    if getattr(item, 'hr_id', None) == current_user.id:
        is_mapped = _resolve_resume_mapping_bool(
            item, db, hr_app_paths, hr_app_emails_no_path, hr_app_emails_job
        )
        return True, is_mapped

    # 2. Visible if mapped to this HR user's job or application
    is_mapped_to_hr = _resolve_resume_mapping_bool(
        item, db, hr_app_paths, hr_app_emails_no_path, hr_app_emails_job
    )
    if is_mapped_to_hr:
        return True, True

    # 3. If mapped globally to another HR, hide it
    is_mapped_globally = _resolve_resume_mapping_bool(
        item, db, global_app_paths, global_app_emails_no_path, global_app_emails_job
    )
    if is_mapped_globally:
        return False, False

    # 4. Visible if unmapped but targets this HR user's job code
    target_job_id = _get_target_job_id_from_subject(item.subject, db)
    if target_job_id and target_job_id in hr_job_ids:
        return True, False

    return False, False


@router.get("/ingested-emails/stats")
def get_ingested_emails_stats(
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """
    Return accurate global counts for the email inbox stats cards:
    total ingested, auto-mapped (has application), and pending assignment.
    """
    items = db.query(AttachmentResume).all()
    
    # Pre-fetch application mapping data for efficient lookups
    applications_data = db.query(Application.resume_file_path, Application.candidate_email, Application.job_id).all()
    
    # Store in sets for O(1) lookups
    global_app_paths = set()
    global_app_emails_no_path = set()
    global_app_emails_job = set()
    
    for a in applications_data:
        if a.candidate_email:
            email_key = a.candidate_email.lower().strip()
            if a.job_id:
                global_app_emails_job.add((email_key, a.job_id))
            path_id = _extract_storage_path_identifier(a.resume_file_path) if a.resume_file_path else None
            if path_id:
                global_app_paths.add(path_id)
            else:
                global_app_emails_no_path.add(email_key)
                
    hr_app_paths = set()
    hr_app_emails_no_path = set()
    hr_app_emails_job = set()
    hr_job_ids = set()
    
    if current_user.role not in ["super_admin", "admin"]:
        hr_job_ids = {j.id for j in db.query(Job.id).filter(Job.hr_id == current_user.id).all()}
        hr_applications_data = db.query(Application.resume_file_path, Application.candidate_email, Application.job_id).outerjoin(Job).filter(
            (Job.hr_id == current_user.id) | (Application.hr_id == current_user.id)
        ).all()
        
        for a in hr_applications_data:
            if a.candidate_email:
                email_key = a.candidate_email.lower().strip()
                if a.job_id:
                    hr_app_emails_job.add((email_key, a.job_id))
                path_id = _extract_storage_path_identifier(a.resume_file_path) if a.resume_file_path else None
                if path_id:
                    hr_app_paths.add(path_id)
                else:
                    hr_app_emails_no_path.add(email_key)

    total_ingested = 0
    auto_mapped = 0
    
    for item in items:
        is_visible, is_mapped = _is_email_visible_to_user(
            item, current_user, db,
            hr_app_paths, hr_app_emails_no_path, hr_app_emails_job,
            global_app_paths, global_app_emails_no_path, global_app_emails_job,
            hr_job_ids
        )
        if is_visible:
            total_ingested += 1
            if is_mapped:
                auto_mapped += 1
                
    return {
        "total_ingested": total_ingested,
        "auto_mapped": auto_mapped,
        "pending_assignment": max(0, total_ingested - auto_mapped),
    }


@router.get("/ingested-emails")
def get_ingested_emails(
    limit: int = 10,
    skip: int = 0,
    search: str = None,
    processed: bool = None,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """
    List all ingested email resumes (HR only).
    Enforces role-based isolation: standard HR users only see emails matching their applications/jobs.
    """
    from app.core.storage import get_signed_url

    query = db.query(AttachmentResume)
    if search:
        query = query.filter(
            or_(
                AttachmentResume.sender_email.ilike(f"%{search}%"),
                AttachmentResume.subject.ilike(f"%{search}%"),
                AttachmentResume.file_name.ilike(f"%{search}%")
            )
        )
    # Fetch all items for accurate mapping-based filtering in Python
    items = query.order_by(AttachmentResume.id.desc()).all()

    # Pre-fetch application mapping data once — O(n) sets for O(1) lookups.
    applications_data = db.query(Application.resume_file_path, Application.candidate_email, Application.job_id).all()
    
    global_app_paths = set()
    global_app_emails_no_path = set()
    global_app_emails_job = set()
    
    for a in applications_data:
        if a.candidate_email:
            email_key = a.candidate_email.lower().strip()
            if a.job_id:
                global_app_emails_job.add((email_key, a.job_id))
            path_id = _extract_storage_path_identifier(a.resume_file_path) if a.resume_file_path else None
            if path_id:
                global_app_paths.add(path_id)
            else:
                global_app_emails_no_path.add(email_key)

    hr_app_paths = set()
    hr_app_emails_no_path = set()
    hr_app_emails_job = set()
    hr_job_ids = set()
    
    if current_user.role not in ["super_admin", "admin"]:
        hr_job_ids = {j.id for j in db.query(Job.id).filter(Job.hr_id == current_user.id).all()}
        hr_applications_data = db.query(Application.resume_file_path, Application.candidate_email, Application.job_id).outerjoin(Job).filter(
            (Job.hr_id == current_user.id) | (Application.hr_id == current_user.id)
        ).all()
        
        for a in hr_applications_data:
            if a.candidate_email:
                email_key = a.candidate_email.lower().strip()
                if a.job_id:
                    hr_app_emails_job.add((email_key, a.job_id))
                path_id = _extract_storage_path_identifier(a.resume_file_path) if a.resume_file_path else None
                if path_id:
                    hr_app_paths.add(path_id)
                else:
                    hr_app_emails_no_path.add(email_key)

    visible_items_with_mapping = []
    for item in items:
        is_visible, is_mapped = _is_email_visible_to_user(
            item, current_user, db,
            hr_app_paths, hr_app_emails_no_path, hr_app_emails_job,
            global_app_paths, global_app_emails_no_path, global_app_emails_job,
            hr_job_ids
        )
        if is_visible:
            visible_items_with_mapping.append((item, is_mapped))

    results = []

    for item, is_mapped in visible_items_with_mapping:
        app_obj = None
        if is_mapped:
            # Only hit the DB when we know a match exists (subset of all rows).
            sender_str = item.sender_email or ""
            match = re.search(r'<([^>]+)>', sender_str)
            raw_email = match.group(1).lower().strip() if match else sender_str.lower().strip()
            path_id = _extract_storage_path_identifier(item.file_url) if item.file_url else None
            
            if path_id:
                app_obj = (
                    db.query(Application)
                    .options(joinedload(Application.job))
                    .filter(
                        or_(
                            Application.resume_file_path == path_id,
                            Application.resume_file_path.like(f"%{path_id}%")
                        )
                    )
                    .first()
                )
            
            if not app_obj and raw_email:
                # Try matching by email where resume_file_path is None (manually mapped no attachment)
                app_obj = (
                    db.query(Application)
                    .options(joinedload(Application.job))
                    .filter(
                        func.trim(func.lower(Application.candidate_email)) == raw_email,
                        or_(
                            Application.resume_file_path == None,
                            Application.resume_file_path == ""
                        )
                    )
                    .first()
                )
            
            if not app_obj and raw_email:
                # Try matching by email and target job ID
                target_job_id = _get_target_job_id_from_subject(item.subject, db)
                if target_job_id:
                    app_obj = (
                        db.query(Application)
                        .options(joinedload(Application.job))
                        .filter(
                            func.trim(func.lower(Application.candidate_email)) == raw_email,
                            Application.job_id == target_job_id
                        )
                        .first()
                    )

        # Duplicate detection
        sender_str = item.sender_email or ""
        match_dup = re.search(r'<([^>]+)>', sender_str)
        raw_email_dup = match_dup.group(1).lower().strip() if match_dup else sender_str.lower().strip()
        candidate_email_to_check = app_obj.candidate_email if app_obj else raw_email_dup

        is_duplicate = False
        if candidate_email_to_check:
            dup_query = db.query(Application).filter(Application.candidate_email.ilike(candidate_email_to_check))
            if app_obj:
                dup_query = dup_query.filter(Application.id != app_obj.id)
            is_duplicate = dup_query.count() > 0

        fresh_file_url = item.file_url
        if item.file_url:
            path_id_for_url = _extract_storage_path_identifier(item.file_url)
            if path_id_for_url:
                storage_key = path_id_for_url
                if storage_key.startswith("MAIL_ATTACHMENTS/"):
                    storage_key = storage_key[len("MAIL_ATTACHMENTS/"):]
                elif storage_key.startswith("resumes/"):
                    storage_key = storage_key[len("resumes/"):]
                try:
                    fresh_url = get_signed_url('MAIL_ATTACHMENTS', storage_key, expires_in=86400)
                    if fresh_url:
                        fresh_file_url = fresh_url
                except Exception as url_err:
                    logger.warning(f"Failed to refresh signed URL for AttachmentResume {item.id}: {url_err}")

        results.append({
            "id": item.id,
            "sender_email": item.sender_email,
            "subject": item.subject,
            "file_name": item.file_name,
            "file_url": fresh_file_url,
            "received_at": item.received_at.replace(tzinfo=timezone.utc) if item.received_at else None,
            "processed": is_mapped or (app_obj is not None),
            "mapping_failed": getattr(item, 'mapping_failed', False),
            "application_id": app_obj.id if app_obj else None,
            "job_title": app_obj.job.title if app_obj and app_obj.job else None,
            "job_code": app_obj.job.job_id if app_obj and app_obj.job else None,
            "is_duplicate": is_duplicate
        })

    # Stats scoping
    if current_user.role not in ["super_admin", "admin"]:
        if search:
            # Stats reflect all items the user is authorized to see (without search)
            all_db_items = db.query(AttachmentResume).order_by(AttachmentResume.id.desc()).all()
            all_visible_items = []
            for item in all_db_items:
                is_visible, is_mapped = _is_email_visible_to_user(
                    item, current_user, db,
                    hr_app_paths, hr_app_emails_no_path, hr_app_emails_job,
                    global_app_paths, global_app_emails_no_path, global_app_emails_job,
                    hr_job_ids
                )
                if is_visible:
                    all_visible_items.append((item, is_mapped))
        else:
            all_visible_items = visible_items_with_mapping

        total_ingested = len(all_visible_items)
        auto_mapped_global = sum(1 for (item, is_mapped) in all_visible_items if is_mapped)
    else:
        total_ingested = db.query(AttachmentResume).count()
        auto_mapped_count = sum(1 for r in results if r["application_id"] is not None)
        
        if search:
            all_items_for_stats = db.query(AttachmentResume).all()
            all_results_app_found = [
                (i, _resolve_resume_mapping(i, db, global_app_paths, global_app_emails_no_path, global_app_emails_job))
                for i in all_items_for_stats
            ]
            auto_mapped_global = sum(1 for (i, (af, ip)) in all_results_app_found if af)
        else:
            auto_mapped_global = auto_mapped_count

    # Python-level filter to match UI's status filter accurately
    if processed is not None:
        if processed:
            results = [r for r in results if r["application_id"] is not None]
        else:
            results = [r for r in results if r["application_id"] is None]

    total = len(results)
    paginated_items = results[skip: skip + limit]

    return {
        "items": paginated_items,
        "total": total,
        "page": (skip // limit) + 1,
        "size": limit,
        "pages": (total + limit - 1) // limit if limit > 0 else 1,
        "global_stats": {
            "total_ingested": total_ingested,
            "auto_mapped": auto_mapped_global,
            "pending_assignment": max(0, total_ingested - auto_mapped_global),
        },
    }




class AssignResumeRequest(BaseModel):
    job_id: int


@router.post("/ingested-emails/{resume_id}/assign")
async def assign_ingested_email(
    resume_id: int,
    req: AssignResumeRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """
    Manually assign an unmapped email resume to a specific job (HR only)
    """
    import re
    from sqlalchemy import or_
    resume = db.query(AttachmentResume).filter(AttachmentResume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Ingested resume not found")

    # Cannot assign an email without a resume attachment — there's nothing to parse
    if not resume.file_url:
        raise HTTPException(
            status_code=400,
            detail="This email has no resume attachment. Only emails with attached resumes (PDF/DOCX) can be assigned to a job."
        )

    # Check if the current user has access to the resume (standard HR visibility check to prevent IDOR)
    if current_user.role not in ["super_admin", "admin"]:
        # 1. Check if the resume was fetched by this HR manager
        if resume.hr_id != current_user.id:
            # 2. Check if the resume is already mapped to one of this HR manager's jobs or applications
            path_id = _extract_storage_path_identifier(resume.file_url) if resume.file_url else None
            sender_str = resume.sender_email or ""
            match_email = re.search(r'<([^>]+)>', sender_str)
            raw_email = match_email.group(1).lower().strip() if match_email else sender_str.lower().strip()
            
            is_mapped_to_hr = False
            if path_id:
                is_mapped_to_hr = db.query(Application).outerjoin(Job).filter(
                    (Job.hr_id == current_user.id) | (Application.hr_id == current_user.id),
                    (Application.resume_file_path == path_id) | (Application.resume_file_path.like(f"%{path_id}%"))
                ).count() > 0
            if not is_mapped_to_hr and raw_email:
                is_mapped_to_hr = db.query(Application).outerjoin(Job).filter(
                    (Job.hr_id == current_user.id) | (Application.hr_id == current_user.id),
                    func.trim(func.lower(Application.candidate_email)) == raw_email,
                    or_(Application.resume_file_path == None, Application.resume_file_path == "")
                ).count() > 0
            
            # 3. Check if the resume is mapped globally to someone else
            is_mapped_globally = False
            if not is_mapped_to_hr:
                if path_id:
                    is_mapped_globally = db.query(Application).filter(
                        (Application.resume_file_path == path_id) | (Application.resume_file_path.like(f"%{path_id}%"))
                    ).count() > 0
                if not is_mapped_globally and raw_email:
                    is_mapped_globally = db.query(Application).filter(
                        func.trim(func.lower(Application.candidate_email)) == raw_email,
                        or_(Application.resume_file_path == None, Application.resume_file_path == "")
                    ).count() > 0

            # 4. Check if it targets this HR user's job code
            is_target_job_owned = False
            if not is_mapped_to_hr and not is_mapped_globally:
                target_job_id = _get_target_job_id_from_subject(resume.subject, db)
                if target_job_id:
                    is_target_job_owned = db.query(Job).filter(
                        Job.id == target_job_id,
                        Job.hr_id == current_user.id
                    ).count() > 0

            if not is_mapped_to_hr and not is_target_job_owned:
                raise HTTPException(status_code=403, detail="Forbidden: You do not have access to this ingested email")
        
    job = db.query(Job).filter(Job.id == req.job_id, Job.status == 'open').first()
    if not job:
        raise HTTPException(status_code=404, detail="Target open job not found")

    # Validate target job ownership for standard HR managers to prevent IDOR
    if current_user.role not in ["super_admin", "admin"] and job.hr_id != current_user.id:
        raise HTTPException(status_code=403, detail="Forbidden: You do not own this job")
        
    # ── Resolve candidate details ──────────────────────────────────
    sender_str = resume.sender_email or ""
    # 1. Try to extract name from "Name <email@addr.com>" format
    name_match = re.search(r'^([^<]+)', sender_str)
    candidate_name = name_match.group(1).strip() if name_match else "Emailed Candidate"
    
    # 2. Extract raw email address
    email_match = re.search(r'<([^>]+)>', sender_str)
    raw_email = email_match.group(1).lower().strip() if email_match else sender_str.lower().strip()
    
    # 3. Handle cases where name is missing or placeholder
    if not candidate_name or candidate_name.lower() == "emailed candidate":
        if raw_email:
            candidate_name = raw_email.split('@')[0].replace('.', ' ').title()
        else:
            candidate_name = "Emailed Candidate"

    # 4. Check for disposable email
    is_disposable = False
    if raw_email:
        try:
            from app.core.config import get_settings
            settings = get_settings()
            if settings.disposable_email_domains:
                DISPOSABLE_DOMAINS = {d.strip().lower() for d in settings.disposable_email_domains.split(",") if d.strip()}
            else:
                DISPOSABLE_DOMAINS = {
                    "fengnu.com", "mailinator.com", "guerrillamail.com", "10minutemail.com", 
                    "tempmail.com", "yopmail.com", "sharklasers.com", "getnada.com", 
                    "dispostable.com", "trashmail.com", "mail.tm", "mail.gw", "temp-mail.org"
                }
            _, domain_part = raw_email.rsplit("@", 1)
            if domain_part.lower().strip() in DISPOSABLE_DOMAINS:
                is_disposable = True
        except Exception:
            pass
        
    # ── Resolve phone details (from body) ──────────────────────────
    body_lower = (resume.email_body or "").lower()
    phone_matches = re.findall(r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]', body_lower)
    candidate_phone_normalized = None
    candidate_phone_hash = None
    candidate_phone_raw = None
    if phone_matches:
        from app.core.phone_utils import normalize_phone_digits, compute_phone_hash
        candidate_phone_raw = phone_matches[0]
        norm_p, _ = normalize_phone_digits(candidate_phone_raw)
        if norm_p and len(norm_p) >= 10:
            candidate_phone_normalized = norm_p
            candidate_phone_hash = compute_phone_hash(norm_p)

    # ── Duplicate check ────────────────────────────────────────────
    filters = []
    if raw_email:
        filters.append(Application.candidate_email.ilike(raw_email))
    if candidate_phone_hash:
        filters.append(Application.candidate_phone_hash == candidate_phone_hash)
        
    if filters:
        existing_app = db.query(Application).filter(
            Application.job_id == job.id,
            or_(*filters)
        ).first()
        if existing_app:
            raise HTTPException(status_code=400, detail="Candidate already has an application for this job")
            
    # ── Storage & Content ──────────────────────────────────────────
    # Resolve the correct bucket and relative path
    resume_file_path = None
    if resume.file_url:
        path_id = _extract_storage_path_identifier(resume.file_url)
        if path_id:
            # Re-prepend bucket if we stripped it, to maintain database consistency
            if path_id.startswith("resumes/"):
                resume_file_path = path_id
            elif not path_id.startswith("MAIL_ATTACHMENTS/"):
                resume_file_path = f"MAIL_ATTACHMENTS/{path_id}"
            else:
                resume_file_path = path_id
        
    # Download file to generate accurate hash
    content = b""
    if resume.file_url:
        from app.core.storage import download_file
        try:
            # Try cloud download first
            content = download_file('MAIL_ATTACHMENTS', resume.file_url)
            
            # Fallback to direct HTTP if it looks like a full URL
            if not content and resume.file_url.startswith("http"):
                from urllib.parse import urlparse
                from app.core.config import get_settings
                settings = get_settings()
                parsed_url = urlparse(resume.file_url)
                if not parsed_url.scheme or parsed_url.scheme.lower() != "https":
                    raise HTTPException(status_code=400, detail="Only HTTPS scheme is allowed for safety.")
                
                # Check netloc/domain
                allowed_domains = []
                if settings.supabase_url:
                    supabase_netloc = urlparse(settings.supabase_url).netloc
                    if supabase_netloc:
                        allowed_domains.append(supabase_netloc)
                
                netloc_lower = parsed_url.netloc.lower()
                is_allowed = netloc_lower.endswith(".supabase.co") or netloc_lower == "supabase.co" or any(d == netloc_lower for d in allowed_domains)
                
                if not is_allowed:
                    logger.error(f"SSRF Prevention: Blocked fetching URL '{resume.file_url}'")
                    raise HTTPException(status_code=400, detail="Domain not allowed for file download.")
                
                import requests
                resp = requests.get(resume.file_url, timeout=10)
                if resp.status_code == 200:
                    content = resp.content
        except Exception as e:
            logger.error(f"Failed to download resume for assignment: {e}")
            
    import hashlib
    resume_hash = hashlib.sha256(content).hexdigest() if content else f"manual_hash_{resume.id}_{int(time.time())}"
    
    # ── Create Application ─────────────────────────────────────────
    logger.info(f"🔗 Manual Assignment: Creating application for resume {resume_id} → job {job.id} ({job.title})")
    logger.info(f"   • Candidate: {candidate_name} <{raw_email}>")
    logger.info(f"   • Resume file: {resume_file_path}")
    
    new_application = Application(
        job_id=job.id,
        hr_id=job.hr_id,
        candidate_name=candidate_name,
        candidate_email=raw_email,
        candidate_phone=candidate_phone_normalized,
        candidate_phone_normalized=None,
        candidate_phone_raw=None,
        candidate_phone_hash=candidate_phone_hash,
        resume_file_name=resume.file_name,
        resume_hash=resume_hash,
        resume_file_path=resume_file_path,
        is_disposable_email=is_disposable,
        status="applied",
        applied_at=get_ist_now(),
        resume_status="pending",
        hr_notes="Manually assigned from Ingested Email Recruiter Channel."
    )
    
    logger.info(f"   • Adding application to database...")
    db.add(new_application)
    
    logger.info(f"   • Updating AttachmentResume {resume_id}: processed=True, mapping_failed=False")
    resume.processed = True
    resume.mapping_failed = False
    
    logger.info(f"   • Committing transaction...")
    db.commit()
    db.refresh(new_application)
    
    logger.info(f"✅ Manual Assignment Complete:")
    logger.info(f"   • Application ID: {new_application.id}")
    logger.info(f"   • Status: {new_application.status}")
    logger.info(f"   • Resume Processed: {resume.processed}")
    logger.info(f"   • Mapping Failed: {resume.mapping_failed}")
    
    # Trigger background AI analysis
    from app.api.applications import process_application_background
    background_tasks.add_task(
        process_application_background,
        new_application.id,
        job.id,
        new_application.resume_file_path,
        raw_email,
        candidate_name
    )
    
    return {
        "status": "success",
        "message": "Resume successfully assigned to job and AI analysis triggered.",
        "application_id": new_application.id
    }

def _get_hr_visibility_sets(current_user: User, db: Session):
    applications_data = db.query(Application.resume_file_path, Application.candidate_email, Application.job_id).all()
    
    global_app_paths = set()
    global_app_emails_no_path = set()
    global_app_emails_job = set()
    
    for a in applications_data:
        if a.candidate_email:
            email_key = a.candidate_email.lower().strip()
            if a.job_id:
                global_app_emails_job.add((email_key, a.job_id))
            path_id = _extract_storage_path_identifier(a.resume_file_path) if a.resume_file_path else None
            if path_id:
                global_app_paths.add(path_id)
            else:
                global_app_emails_no_path.add(email_key)

    hr_app_paths = set()
    hr_app_emails_no_path = set()
    hr_app_emails_job = set()
    hr_job_ids = set()
    
    if current_user.role not in ["super_admin", "admin"]:
        hr_job_ids = {j.id for j in db.query(Job.id).filter(Job.hr_id == current_user.id).all()}
        hr_applications_data = db.query(Application.resume_file_path, Application.candidate_email, Application.job_id).outerjoin(Job).filter(
            (Job.hr_id == current_user.id) | (Application.hr_id == current_user.id)
        ).all()
        
        for a in hr_applications_data:
            if a.candidate_email:
                email_key = a.candidate_email.lower().strip()
                if a.job_id:
                    hr_app_emails_job.add((email_key, a.job_id))
                path_id = _extract_storage_path_identifier(a.resume_file_path) if a.resume_file_path else None
                if path_id:
                    hr_app_paths.add(path_id)
                else:
                    hr_app_emails_no_path.add(email_key)
                    
    return (
        hr_app_paths,
        hr_app_emails_no_path,
        hr_app_emails_job,
        global_app_paths,
        global_app_emails_no_path,
        global_app_emails_job,
        hr_job_ids,
    )

@router.delete("/ingested-emails/{resume_id}")
def delete_ingested_email(
    resume_id: int,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """
    Delete an ingested email record (HR / Admin only).
    """
    if current_user.role not in ["super_admin", "admin", "hr"]:
        raise HTTPException(status_code=403, detail="Forbidden: Insufficient permissions to delete ingested emails")

    item = db.query(AttachmentResume).filter(AttachmentResume.id == resume_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Ingested resume not found")

    # Check if already assigned (mapped)
    app_found, _ = _resolve_resume_mapping(item, db)
    if app_found:
        raise HTTPException(
            status_code=400,
            detail="Cannot delete an ingested email that is already assigned to a job."
        )

    # For non-admins, check visibility
    if current_user.role not in ["super_admin", "admin"]:
        hr_sets = _get_hr_visibility_sets(current_user, db)
        is_visible, _ = _is_email_visible_to_user(
            item, current_user, db,
            *hr_sets
        )
        if not is_visible:
            raise HTTPException(status_code=403, detail="Forbidden: Insufficient permissions to delete this ingested email.")
        
    db.delete(item)
    db.commit()
    return {"status": "success", "message": "Ingested email deleted successfully."}

@router.post("/ingested-emails/bulk-delete")
def bulk_delete_ingested_emails(
    payload: BulkDeleteEmailsRequest,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """
    Bulk delete ingested email records (HR / Admin only).
    """
    if current_user.role not in ["super_admin", "admin", "hr"]:
        raise HTTPException(status_code=403, detail="Forbidden: Insufficient permissions to delete ingested emails")

    if not payload.ids:
        return {"status": "success", "message": "No IDs provided.", "deleted_count": 0}

    # Fetch items
    items = db.query(AttachmentResume).filter(AttachmentResume.id.in_(payload.ids)).all()
    if len(items) != len(set(payload.ids)):
        found_ids = {item.id for item in items}
        missing_ids = set(payload.ids) - found_ids
        if missing_ids:
            raise HTTPException(
                status_code=404,
                detail=f"Some ingested email records were not found: {list(missing_ids)}"
            )

    # Check mapping and visibility
    hr_sets = None
    if current_user.role not in ["super_admin", "admin"]:
        hr_sets = _get_hr_visibility_sets(current_user, db)

    for item in items:
        # Check if already assigned (mapped)
        app_found, _ = _resolve_resume_mapping(item, db)
        if app_found:
            raise HTTPException(
                status_code=400,
                detail=f"Cannot delete ingested email ID {item.id} because it is already assigned to a job."
            )

        # For non-admins, check visibility
        if hr_sets:
            is_visible, _ = _is_email_visible_to_user(
                item, current_user, db,
                *hr_sets
            )
            if not is_visible:
                raise HTTPException(
                    status_code=403,
                    detail=f"Forbidden: Insufficient permissions to delete ingested email ID {item.id}."
                )

    # All checks passed, perform deletion
    deleted_count = 0
    for item in items:
        db.delete(item)
        deleted_count += 1

    db.commit()
    return {"status": "success", "message": f"Successfully deleted {deleted_count} ingested emails.", "deleted_count": deleted_count}


@router.get("/{application_id}", response_model=ApplicationDetailResponse)
def get_application(
    application_id: int,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Get application details (HR only)"""
    application = db.query(Application).options(
        joinedload(Application.job),
        joinedload(Application.hr),
        joinedload(Application.resume_extraction),
        joinedload(Application.interview),
        selectinload(Application.pipeline_stages)
    ).filter(Application.id == application_id).first()
    
    if not application:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Application not found"
        )
    validate_hr_ownership(application, current_user, resource_name="application")
    
    # Automatic Self-Healing for Incompatible/Mismatched Encryption Keys
    if application.resume_extraction and application.resume_file_path:
        ext_text = application.resume_extraction.extracted_text
        if ext_text in ("[UNREADABLE]", "[DECRYPTION_ERROR]") or not ext_text or ext_text.strip() == "":
            logger.warning(f"Decryption mismatch or missing text detected for Application ID {application_id}. Triggering automatic self-healing re-parse.")
            
            application.resume_status = "parsing"
            db.commit()
            db.refresh(application)
            
            background_tasks.add_task(
                retry_application_background,
                application_id,
                application.job_id,
                application.resume_file_path
            )
    
    # Sanitize paths using pathlib.Path.resolve() and a safe root check
    def sanitize_path(path_str):
        if not path_str or not (":" in path_str or "\\" in path_str or ".." in path_str):
            return path_str
        try:
            from pathlib import Path
            base_dir = Path("uploads").resolve()
            idx = path_str.find("uploads")
            if idx != -1:
                rel_path = path_str[idx + len("uploads"):].lstrip("\\/")
                full_path = (base_dir / rel_path).resolve()
            else:
                full_path = Path(path_str).resolve()
            
            # Safe root check
            if base_dir in full_path.parents or full_path == base_dir:
                return "uploads/" + full_path.relative_to(base_dir).as_posix()
            else:
                return full_path.name
        except Exception:
            return path_str.replace("\\", "/").split("/")[-1]

    application.candidate_photo_path = sanitize_path(application.candidate_photo_path)
    application.resume_file_path = sanitize_path(application.resume_file_path)

    return build_application_detail_response(application, current_user.id)

async def retry_application_background(application_id: int, job_id: int, bucket_path: str):
    """Safely retry AI resume extraction without altering pipeline stages or triggering emails."""
    db = SessionLocal()
    try:
        from app.services.candidate_service import CandidateService
        cand_service = CandidateService(db)
        
        # Reload objects in this session
        # Step 1: lock only (without joins)
        # Note: Application.resume_extraction is configured as lazy='joined', so ORM-level
        # .with_for_update() may turn into a LEFT OUTER JOIN, which Postgres rejects.
        db.execute(text("SELECT 1 FROM applications WHERE id = :id FOR UPDATE"), {"id": application_id})
        
        # Step 2: fetch with joins, no lock
        application = db.query(Application).options(
            joinedload(Application.resume_extraction)
        ).filter(Application.id == application_id).first()
        
        job = db.query(Job).filter(Job.id == job_id).first()
        if not application or not job:
            db.close()
            return
            
        logger.info(f"Retrying AI extraction for application {application_id}...")
        
        # Parse resume text based on file type from Supabase
        resume_text = ""
        try:
            from io import BytesIO
            from app.core.storage import download_file
            response = download_file(settings.supabase_bucket_resumes, bucket_path)
            if not response:
                raise Exception("Failed to download resume content")
            file_stream = BytesIO(response)
            file_ext = bucket_path.lower().split('_')[-1].split('.')[-1] if '.' in bucket_path else 'pdf'

            if file_ext == 'pdf':
                from pypdf import PdfReader
                reader = PdfReader(file_stream)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        resume_text += page_text + "\n"
            elif file_ext == 'docx':
                import docx
                doc = docx.Document(file_stream)
                for para in doc.paragraphs:
                    if para.text:
                        resume_text += para.text + "\n"
            else:
                file_stream.seek(0)
                resume_text = file_stream.read().decode('utf-8', errors='ignore')

            # Post-extraction sanity check
            if file_ext == 'pdf' and len(response) > 50000 and len(resume_text.strip()) < 100:
                resume_text = "[[SCANNED_PDF_DETECTED]]\n" + resume_text
        except Exception as e:
            logger.error(f"Retry Text Extraction Error: {e}", exc_info=True)
            cand_service.create_audit_log(None, "RETRY_TEXT_EXTRACTION_FAILED", "Application", application_id, {"error": str(e)})
            resume_text = "Error extracting text."
        
        if not resume_text.strip():
            resume_text = "No readable text found."

        # AI Parsing
        extraction_data = await parse_resume_with_ai(resume_text, job_id, job.description, job.experience_level)
        extraction_degraded_flag = extraction_data.pop("extraction_degraded", False)

        # Look for existing extraction or create new
        resume_extraction = db.query(ResumeExtraction).filter(ResumeExtraction.application_id == application_id).first()
        if not resume_extraction:
            resume_extraction = ResumeExtraction(application_id=application_id)
            db.add(resume_extraction)
            
        resume_extraction.extracted_text = resume_text
        resume_extraction.summary = extraction_data.get("summary", "")
        resume_extraction.extracted_skills = json.dumps(extraction_data.get("skills") or [])
        resume_extraction.years_of_experience = extraction_data.get("experience")
        resume_extraction.education = json.dumps(extraction_data.get("education") or [])
        resume_extraction.previous_roles = json.dumps(extraction_data.get("roles") or [])
        resume_extraction.experience_level = extraction_data.get("experience_level")
        resume_extraction.resume_score = extraction_data.get("score", 0)
        resume_extraction.skill_match_percentage = extraction_data.get("match_percentage", 0)
        
        if extraction_data.get("candidate_name"):
            resume_extraction.candidate_name = extraction_data.get("candidate_name")
        if extraction_data.get("email"):
            resume_extraction.email = extraction_data.get("email")
        if extraction_data.get("phone_number"):
            resume_extraction.phone_number = extraction_data.get("phone_number")
        
        resume_extraction.reasoning = {"ai_justification": extraction_data.get("reasoning")}
        
        application.resume_score = extraction_data.get("score", 0)
        application.scoring_metadata = {
            "logic_version": "v2.0",
            "weights": {"skills": 0.6, "experience": 0.4},
            "recomputed_at": get_ist_now().isoformat(),
            "extraction_degraded": extraction_degraded_flag
        }
        
        # ── HYBRID IDENTITY EXTRACTION (Sync on retry) ──
        from app.services.ai_service import extract_email_regex, extract_phone_regex, extract_name_heuristic
        extracted_name = extraction_data.get("candidate_name") or extract_name_heuristic(resume_text)
        extracted_email = extraction_data.get("email") or extract_email_regex(resume_text)
        extracted_phone = extraction_data.get("phone_number") or extract_phone_regex(resume_text)
        
        email_is_placeholder = application.candidate_email and "@batch." in application.candidate_email
        name_is_placeholder = not application.candidate_name or len(application.candidate_name.split()) < 2

        if extracted_email and (not application.candidate_email or email_is_placeholder):
            application.candidate_email = extracted_email
        if extracted_name and (not application.candidate_name or name_is_placeholder):
            application.candidate_name = extracted_name
        if extracted_phone and not application.candidate_phone:
            application.candidate_phone = extracted_phone

        if extraction_degraded_flag:
            _append_extraction_degraded_marker(application)

        application.resume_status = "parsed"; application.failure_reason = None
        
        # ── Pipeline Advancement on Retry ──
        # Same logic as initial: pass unless it's a known duplicate (retry usually shouldn't happen for duplicates)
        stage_status = "pass"
        stage_note = "AI analysis complete (via manual retry) — awaiting HR decision"
        
        cand_service.advance_stage(
            application_id, 
            "Resume Screening", 
            stage_status, 
            extraction_data.get("score", 0) * 10, 
            stage_note
        )

        db.commit()
        cand_service.create_audit_log(None, "AI_ANALYSIS_RETRY_SUCCESS", "Application", application_id, {"score": extraction_data.get("score")})
        logger.info(f"Retry successful for application {application_id}")
        try:
            log_json(
                logger,
                "resume_retry_background_completed",
                level="info",
                extra={"application_id": application_id, "resume_status": "parsed"},
            )
        except Exception:
            pass
    except Exception as e:
        logger.error(
            f"Retry Background Error processing application {application_id}: {e}",
            exc_info=True,
        )
        db.rollback()
        try:
            cand_service = CandidateService(db)
            cand_service.create_audit_log(None, "AI_ANALYSIS_RETRY_FAILED", "Application", application_id, {"error": str(e)})
            failed_app = db.query(Application).filter(Application.id == application_id).first()
            if failed_app:
                failed_app.resume_status = "failed"
                failed_app.retry_count = (failed_app.retry_count or 0) + 1
                failed_app.failure_reason = str(e)[:1000] # Cap length
                # Error details are now handled by the frontend via failure_reason
                db.commit()
            try:
                log_json(
                    logger,
                    "resume_retry_background_failed",
                    level="error",
                    extra={"application_id": application_id, "resume_status": "failed"},
                )
            except Exception:
                pass
        except Exception:
            pass
    finally:
        db.close()

@router.post("/{application_id}/retry-analysis")
async def retry_resume_analysis(
    application_id: int, 
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Manually trigger AI resume analysis if it failed"""
    application = db.query(Application).filter(Application.id == application_id).with_for_update().first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
    validate_hr_ownership(application, current_user, resource_name="application")
        
    application.resume_status = "parsing"
    application.updated_at = get_ist_now()
    db.commit()
    db.refresh(application)
    try:
        log_json(
            logger,
            "resume_retry_api_accepted",
            level="info",
            extra={
                "application_id": application_id,
                "hr_user_id": current_user.id,
                "resume_status": "parsing",
            },
        )
    except Exception:
        pass

    background_tasks.add_task(
        retry_application_background, 
        application.id, 
        application.job_id, 
        application.resume_file_path
    )
    
    return {
        "status": "success",
        "message": "Analysis restarted in background safely",
        "application_id": application_id,
        "resume_status": "parsing",
    }

@router.post("/{application_id}/resend-interview-invitation")
async def resend_interview_invitation(
    application_id: int,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db),
):
    """
    Re-send the interview invitation email.

    This is useful when the original "approve_for_interview" succeeded but the email
    provider (e.g. Gmail quota) failed. We re-issue the interview access key only when
    the interview is missing or still 'not_started'.
    """
    # Concurrency Hardening (Phase 1): Lock row for update
    application = db.query(Application).options(
        joinedload(Application.job),
        joinedload(Application.interview),
    ).filter(Application.id == application_id).with_for_update().first()

    if not application:
        raise HTTPException(status_code=404, detail="Application not found")

    validate_hr_ownership(application, current_user, resource_name="application")

    # If an interview exists and is already in progress/completed, don't re-issue the key.
    if application.interview and getattr(application.interview, "status", None) != "not_started":
        raise HTTPException(
            status_code=400,
            detail="Interview access key cannot be reissued unless interview is 'not_started'",
        )

    from app.services.candidate_service import CandidateService
    cand_service = CandidateService(db)
    raw_access_key = cand_service.ensure_interview_record_exists(application)
    candidate_email = application.candidate_email
    job_title = application.job.title if application.job else "your applied position"

    background_tasks.add_task(
        send_approved_for_interview_email,
        candidate_email,
        job_title,
        raw_access_key,
    )

    try:
        log_json(
            logger,
            "resume_invite_email_resend_scheduled",
            level="info",
            extra={"application_id": application_id, "to": candidate_email},
        )
    except Exception:
        pass

    return {
        "status": "success",
        "message": "Interview invitation email scheduled in background",
        "application_id": application_id,
    }

@router.put("/{application_id}/status")
async def update_application_status(
    application_id: int,
    status_update: ApplicationStatusUpdate,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """
    Execute a state transition on an application via the finite state machine.
    
    The frontend sends an 'action' (e.g. 'approve_for_interview', 'reject', 
    'call_for_interview', 'review_later', 'hire').  The FSM validates the 
    transition, updates the status atomically, logs the change, and returns 
    the result including which email to send.
    """
    from app.services.state_machine import (
        CandidateStateMachine, TransitionAction,
        InvalidTransitionError, DuplicateTransitionError,
        get_user_friendly_fsm_error,
    )

    application = db.query(Application).filter(Application.id == application_id).with_for_update().first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
    validate_hr_ownership(application, current_user, resource_name="application")

    # Parse the action
    try:
        action = TransitionAction(status_update.action)
    except ValueError:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid action: '{status_update.action}'. "
                   f"Valid actions: {[a.value for a in TransitionAction if not a.value.startswith('system_')]}"
        )

    # Block system actions from HR endpoint
    if action.value.startswith("system_"):
        raise HTTPException(status_code=400, detail="System actions cannot be triggered manually")

    # Execute FSM transition
    old_state = application.status
    fsm = CandidateStateMachine(db)
    try:
        logger.debug(f"/api/auth/me user_id={current_user.id}, role={current_user.role}")
        # Hardening Phase 5: Single transaction for Status + Audit Log
        result = fsm.transition(
            application=application,
            action=action,
            user_id=current_user.id,
            notes=status_update.hr_notes,
            is_critical=True,
            background_tasks=background_tasks
        )
        # Flush to confirm DB state without committing yet
        db.flush()
    except (InvalidTransitionError, DuplicateTransitionError) as e:
        raise HTTPException(status_code=400, detail=get_user_friendly_fsm_error(e))

    # Handle HR notes
    if status_update.hr_notes:
        application.hr_notes = status_update.hr_notes

    # ─── Post-transition side effects ───────────────────────────────────
    raw_access_key = None

    if action == TransitionAction.APPROVE_FOR_INTERVIEW:
        # Create or refresh interview record + access key
        from app.services.candidate_service import CandidateService
        cand_service = CandidateService(db)
        raw_access_key = cand_service.ensure_interview_record_exists(application)

    if action == TransitionAction.HIRE:
        # Create HiringDecision record
        from app.domain.models import HiringDecision
        existing = db.query(HiringDecision).filter(
            HiringDecision.application_id == application_id
        ).first()
        if not existing:
            decision = HiringDecision(
                application_id=application_id,
                hr_id=current_user.id,
                decision="hired",
                decision_comments=status_update.hr_notes or "Hired via pipeline",
                decided_at=get_ist_now(),
            )
            db.add(decision)

    if action in (TransitionAction.HIRE, TransitionAction.ACCEPT_OFFER):
        # Auto-withdraw other active duplicate applications of the same candidate
        candidate_email = application.candidate_email
        if candidate_email:
            from app.domain.constants import CandidateState
            active_duplicates = db.query(Application).filter(
                Application.candidate_email == candidate_email,
                Application.id != application.id,
                Application.status.notin_([
                    CandidateState.REJECTED.value,
                    CandidateState.ONBOARDED.value
                ])
            ).all()
            
            for dup in active_duplicates:
                try:
                    fsm.transition(
                        application=dup,
                        action=TransitionAction.REJECT,
                        user_id=current_user.id,
                        notes=f"Auto-withdrawn: Candidate accepted another offer (App ID: {application.id})",
                        is_critical=False,
                        background_tasks=background_tasks
                    )
                except Exception as e:
                    logger.error(f"Failed to auto-withdraw duplicate App {dup.id}: {e}")

    if action == TransitionAction.REJECT:
        from app.domain.models import HiringDecision, Notification
        existing = db.query(HiringDecision).filter(
            HiringDecision.application_id == application_id
        ).first()
        if not existing:
            decision = HiringDecision(
                application_id=application_id,
                hr_id=current_user.id,
                decision="rejected",
                decision_comments=status_update.hr_notes or "Rejected via pipeline",
                decided_at=get_ist_now(),
            )
            db.add(decision)
        
        # Notify HR if offer was rejected during approval phase
        if old_state == "pending_approval" and application.hr_id:
            db.add(Notification(
                user_id=application.hr_id,
                notification_type="OFFER_REJECTED",
                title="Offer Request Rejected",
                message=f"The offer request for {application.candidate_name} was rejected by Super Admin.",
                related_application_id=application.id
            ))
    # ────────────────────────────────────────────────────────────────────

    # Atomic commit
    try:
        db.commit()
        db.refresh(application)
    except Exception as e:
        db.rollback()
        raise HTTPException(status_code=500, detail="Failed to save state transition")

    # ─── Email triggers (ONLY after successful commit) ──────────────────
    candidate_email = application.candidate_email
    job_title = application.job.title

    logger.info(
        f"EMAIL_TRIGGER_CHECK: app={application.id}, "
        f"email_type={result.email_type}, "
        f"raw_access_key_exists={raw_access_key is not None}, "
        f"candidate_email={candidate_email}"
    )

    if result.email_type == "approved_for_interview" and raw_access_key:
        logger.info(f"[EMAIL] Scheduling interview invitation email to {candidate_email}")
        from app.services.email_service import send_interview_invitation_email
        background_tasks.add_task(send_interview_invitation_email, application, raw_access_key)
    elif result.email_type == "screened":
        logger.info(f"[EMAIL] Scheduling screened email to {candidate_email}")
        from app.services.email_service import send_screened_email
        background_tasks.add_task(send_screened_email, candidate_email, job_title, application)
    elif result.email_type == "rejected":
        logger.info(f"[EMAIL] Scheduling rejected email to {candidate_email}")
        background_tasks.add_task(send_rejected_email, candidate_email, job_title, False, application)
    elif result.email_type == "call_for_interview":
        logger.info(f"[EMAIL] Scheduling call_for_interview email to {candidate_email}")
        from app.services.email_service import send_call_for_interview_email
        background_tasks.add_task(send_call_for_interview_email, candidate_email, job_title)
    elif result.email_type == "hired":
        logger.info(f"[EMAIL] Scheduling hired email to {candidate_email}")
        from app.services.email_service import send_hired_email
        background_tasks.add_task(send_hired_email, candidate_email, job_title, application.interview, None, application)
    elif result.email_type:
        logger.warning(f"[EMAIL] No email trigger matched for email_type={result.email_type}")
    # ────────────────────────────────────────────────────────────────────

    return {
        "id": application.id,
        "status": application.status,
        "transition": {
            "from_state": result.from_state,
            "to_state": result.to_state,
            "action": result.action,
            "email_type": result.email_type,
        }
    }



@router.delete("/{application_id}")
async def delete_application(
    application_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_hr)
):
    """
    Delete an application along with associated data. HR only.
    """
    app = db.query(Application).filter(Application.id == application_id).first()
    if not app:
        raise HTTPException(status_code=404, detail="Application not found")
    validate_hr_ownership(app, current_user, resource_name="application")

    # Fetch file paths before deletion
    resume_path = app.resume_file_path
    photo_path = app.candidate_photo_path
    id_card_path = getattr(app, 'id_card_url', None)
    video_path = app.interview.video_recording_path if app.interview else None

    try:
        # Explicitly delete ResumeExtraction to prevent ForeignKeyViolation
        # (Table missing ON DELETE CASCADE and relationship has passive_deletes=True)
        from app.domain.models import ResumeExtraction
        db.query(ResumeExtraction).filter(ResumeExtraction.application_id == application_id).delete()
        
        db.delete(app)
        db.commit()

        # Delete cloud files after successful commit
        from app.core.storage import delete_file
        if resume_path:
            delete_file(settings.supabase_bucket_resumes, resume_path)
        if photo_path:
            delete_file(settings.supabase_bucket_id_photos, photo_path)
        if id_card_path:
            delete_file(settings.supabase_bucket_id_cards, id_card_path)
        if video_path:
            delete_file(settings.supabase_bucket_videos, video_path)

    except Exception as e:
        db.rollback()
        logger.error(f"Error deleting application {application_id}: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete application. It might have complex dependencies.")
    return {"message": "Application deleted successfully"}
    
@router.post("/bulk-delete")
async def bulk_delete_applications(
    application_ids: List[int],
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_hr)
):
    """
    Delete multiple applications in a single transaction. HR only.
    """
    if not application_ids:
        return {"message": "No applications provided for deletion"}

    if len(application_ids) > 100:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Cannot delete more than 100 applications at a time."
        )

    # Fetch all applications and validate ownership
    apps = db.query(Application).filter(Application.id.in_(application_ids)).all()
    
    if len(apps) != len(application_ids):
        found_ids = [app.id for app in apps]
        missing_ids = list(set(application_ids) - set(found_ids))
        logger.warning(f"Bulk delete: Some applications not found: {missing_ids}")

    for app in apps:
        validate_hr_ownership(app, current_user, resource_name="application")

    # Collect all file paths before deletion
    file_deletions = []
    for app in apps:
        if app.resume_file_path:
            file_deletions.append((settings.supabase_bucket_resumes, app.resume_file_path))
        if app.candidate_photo_path:
            file_deletions.append((settings.supabase_bucket_id_photos, app.candidate_photo_path))
        id_card = getattr(app, 'id_card_url', None)
        if id_card:
            file_deletions.append((settings.supabase_bucket_id_cards, id_card))
        if app.interview and app.interview.video_recording_path:
            file_deletions.append((settings.supabase_bucket_videos, app.interview.video_recording_path))

    try:
        from app.domain.models import ResumeExtraction
        # Delete associated data first
        db.query(ResumeExtraction).filter(ResumeExtraction.application_id.in_(application_ids)).delete(synchronize_session=False)
        
        # Delete the applications
        db.query(Application).filter(Application.id.in_(application_ids)).delete(synchronize_session=False)
        
        db.commit()

        # Delete cloud files after successful commit
        from app.core.storage import delete_file
        for bucket, path in file_deletions:
            try:
                delete_file(bucket, path)
            except Exception as f_err:
                logger.warning(f"Bulk delete: failed to remove cloud file {bucket}/{path}: {f_err}")

    except Exception as e:
        db.rollback()
        logger.error(f"Error in bulk delete: {e}")
        raise HTTPException(status_code=500, detail="Failed to delete applications. Some might have complex dependencies.")
        
    return {"message": f"Successfully deleted {len(apps)} applications"}

@router.post("/{application_id}/merge/{target_id}")
async def merge_applications(
    application_id: int,
    target_id: int,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """
    Merge a source application into a target application.
    Used for resolving duplicate submissions.
    """
    source = db.query(Application).filter(Application.id == application_id).first()
    target = db.query(Application).filter(Application.id == target_id).first()
    
    if not source or not target:
        raise HTTPException(status_code=404, detail="One or both applications not found")
    validate_hr_ownership(source, current_user, resource_name="application")
    validate_hr_ownership(target, current_user, resource_name="application")
        
    if source.job_id != target.job_id:
        raise HTTPException(status_code=400, detail="Applications must belong to the same job to be merged")
        
    # Merge strategy: Target keeps its identity, but takes scores/notes from source if they are better
    try:
        if source.resume_score > (target.resume_score or 0):
            target.resume_score = source.resume_score
            
        target.hr_notes = (target.hr_notes or "") + f"\n[MERGED from App #{application_id}]: " + (source.hr_notes or "No notes")
        
        # Log the merge
        from app.domain.models import AuditLog
        merge_log = AuditLog(
            user_id=current_user.id,
            action="APPLICATION_MERGED",
            resource_type="Application",
            resource_id=target.id,
            details=json.dumps({"source_id": application_id, "target_id": target_id})
        )
        db.add(merge_log)
        
        # Mark source as rejected/duplicate and hide it or delete it
        source.status = "rejected"
        source.hr_notes = (source.hr_notes or "") + f"\n[MERGED INTO App #{target_id}]"
        
        db.commit()
    except Exception as e:
        db.rollback()
        logger.error(f"Merge error: {e}")
        raise HTTPException(status_code=500, detail="Failed to merge applications")
        
    return {"status": "success", "message": f"Application {application_id} merged into {target_id}"}
@router.put("/{application_id}/notes", response_model=ApplicationResponse)
async def update_hr_notes(
    application_id: int,
    notes_update: ApplicationNotesUpdate,
    current_user: User = Depends(get_current_hr),
    db: Session = Depends(get_db)
):
    """Update HR notes for an application"""
    application = db.query(Application).filter(Application.id == application_id).with_for_update().first()
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")
    validate_hr_ownership(application, current_user, resource_name="application")
        
    application.hr_notes = notes_update.hr_notes
    db.commit()
    db.refresh(application)
    return application

@router.post("/extract-basic-info")
@limiter.limit("5/minute")
async def extract_basic_info(
    request: Request,
    resume_file: UploadFile = File(...)
):
    """Fast endpoint to extract Name and Phone from an uploaded resume."""
    # 1. Extension Check
    allowed_extensions = {'.pdf', '.docx', '.doc'}
    file_ext = os.path.splitext(resume_file.filename)[1].lower()
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_ext}")

    # 2. Size Limit Check (2MB cap)
    MAX_FILE_SIZE = 2 * 1024 * 1024
    content = await resume_file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File size exceeds the 2MB limit.")
    
    # Extract text locally in memory
    resume_text = ""
    file_ext = resume_file.filename.lower().split('.')[-1]
    
    try:
        if file_ext == 'pdf':
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(content))
            for page in reader.pages:
                resume_text += page.extract_text() + "\n"
        elif file_ext in ['docx', 'doc']:
            import docx
            import io
            doc = docx.Document(io.BytesIO(content))
            for para in doc.paragraphs:
                resume_text += para.text + "\n"
        else:
            resume_text = content.decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"Error extracting text for basic info: {e}")
        return {"name": "", "phone": ""}
        
    if not resume_text.strip():
        return {"name": "", "phone": ""}
        
    info = await extract_basic_candidate_info(resume_text)
    # Privacy & correctness: never return or pre-fill email here.
    # Only expose minimal fields required for UX pre-fill.
    return {
        "name": info.get("name") if isinstance(info, dict) else "",
        "phone": info.get("phone") if isinstance(info, dict) else "",
    }

